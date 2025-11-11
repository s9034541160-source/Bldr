"""
Генерация документов предварительного ТЭО.
"""

from __future__ import annotations

import logging
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Optional
import tempfile

from docx import Document
from docx2pdf import convert
from jinja2 import Environment, FileSystemLoader, Template
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font

from backend.models.project import Project

logger = logging.getLogger(__name__)


class TEOReportService:
    """
    Формирует отчёты по предварительному ТЭО (DOCX, XLSX, PDF).
    """

    def __init__(self) -> None:
        templates_dir = Path(__file__).resolve().parent.parent / "templates"
        self._env = Environment(loader=FileSystemLoader(templates_dir), autoescape=False)

    def build_preliminary_report(self, project: Project, analysis: Optional[Dict[str, Any]]) -> bytes:
        document = Document()

        document.add_heading("Предварительное ТЭО", level=1)
        document.add_paragraph(f"Проект: {project.name} ({project.code})")
        document.add_paragraph(f"Дата формирования: {datetime.utcnow():%d.%m.%Y %H:%M}")

        summary_text = self._render_text_summary(project, analysis or {})
        if summary_text:
            for line in summary_text.splitlines():
                stripped = line.strip()
                if stripped:
                    document.add_paragraph(stripped)

        document.add_heading("Общая информация", level=2)
        table = document.add_table(rows=0, cols=2)
        rows = [
            ("UUID", project.uuid),
            ("Статус", project.status),
            ("Ожидаемое начало", project.expected_start or ""),
            ("Ожидаемое завершение", project.expected_completion or ""),
            ("Плановая длительность (дн.)", project.planned_duration_days or ""),
        ]
        for title, value in rows:
            cells = table.add_row().cells
            cells[0].text = str(title)
            cells[1].text = str(value)

        if analysis:
            work_volume = analysis.get("work_volume")
            if isinstance(work_volume, dict):
                document.add_heading("Объёмы работ", level=2)
                summary = work_volume.get("summary") or {}
                document.add_paragraph(f"Позиции: {summary.get('total_items', 0)}")
                document.add_paragraph(
                    f"Суммарный объём: {summary.get('total_quantity', 0)} {self._detect_unit(work_volume)}"
                )

            cost = analysis.get("cost")
            if isinstance(cost, dict):
                document.add_heading("Стоимость", level=2)
                base = cost.get("total_cost", 0)
                overhead = cost.get("overhead_cost", 0)
                profit = cost.get("profit", 0)
                grand_total = cost.get("grand_total", base)
                total_with_travel = cost.get("total_with_travel")
                margin = cost.get("margin_percent", 0)
                document.add_paragraph(f"Базовая стоимость: {base:,.2f} руб.")
                document.add_paragraph(f"Накладные расходы (18%): {overhead:,.2f} руб.")
                document.add_paragraph(f"Прибыль (10%): {profit:,.2f} руб.")
                document.add_paragraph(f"Итого с накладными и прибылью: {grand_total:,.2f} руб.")
                if total_with_travel:
                    document.add_paragraph(f"Итого с командировочными и СИЗ: {total_with_travel:,.2f} руб.")
                document.add_paragraph(f"Маржа: {margin:.2f}%")
                breakdown = cost.get("group_breakdown") or {}
                if breakdown:
                    document.add_paragraph("Структура стоимости:")
                    for label, value in breakdown.items():
                        document.add_paragraph(f"- {label.title()}: {value:,.2f} руб.")
                missing = cost.get("missing_prices") or []
                if missing:
                    document.add_paragraph("Позиции без цены:", style="List Bullet")
                    for item in missing:
                        document.add_paragraph(str(item), style="List Bullet")

            labor = analysis.get("labor")
            if isinstance(labor, dict):
                document.add_heading("Трудозатраты", level=2)
                document.add_paragraph(f"Трудоёмкость: {labor.get('total_labor_hours', 0):,.2f} чел·ч")
                document.add_paragraph(
                    f"Эквивалент работников (смена 8ч): {labor.get('total_worker_equivalent', 0):,.2f}"
                )
                schedules = labor.get("schedules") or {}
                if schedules:
                    document.add_paragraph("Рекомендации по вахтовым графикам:")
                    for name, data in schedules.items():
                        document.add_paragraph(
                            f"- {name} → требуемо бригад: {data.get('required_workers', 0)} "
                            f"(цикл {data.get('cycle_days')} дн., работа {data.get('work_days')} дн.)"
                        )

            travel = analysis.get("travel")
            if isinstance(travel, dict):
                document.add_heading("Командировочные и СИЗ", level=2)
                document.add_paragraph(f"Командировочные расходы: {travel.get('total_travel', 0):,.2f} руб.")
                document.add_paragraph(f"Билеты: {travel.get('tickets', 0):,.2f} руб.")
                document.add_paragraph(f"Проживание: {travel.get('lodging', 0):,.2f} руб.")
                document.add_paragraph(f"Суточные: {travel.get('per_diem', 0):,.2f} руб.")
                document.add_paragraph(f"СИЗ: {travel.get('ppe', 0):,.2f} руб.")
                document.add_paragraph(
                    f"Сумма до коэффициентов: {travel.get('total_before_coeff', 0):,.2f} руб."
                )
                document.add_paragraph(
                    f"Итого с коэффициентами: {travel.get('total_with_coeff', 0):,.2f} руб. "
                    f"(коэффициент {travel.get('coefficient', 1.0)})"
                )

            financial = analysis.get("financial")
            if isinstance(financial, dict):
                document.add_heading("Финансовые показатели", level=2)
                document.add_paragraph(f"NPV: {financial.get('npv', 0):,.2f} руб.")
                irr_value = financial.get("irr_percent")
                if irr_value is not None:
                    document.add_paragraph(f"IRR: {irr_value:.2f}%")
                payback = financial.get("payback_months")
                if payback is not None:
                    document.add_paragraph(f"Срок окупаемости: {payback} мес.")
                risk = financial.get("risk") or {}
                document.add_paragraph(
                    f"Уровень риска: {risk.get('level', 'n/a')} (score={risk.get('score', 0):.0f})"
                )
                for factor in risk.get("factors", []):
                    document.add_paragraph(factor, style="List Bullet")

            timeline = analysis.get("timeline")
            if isinstance(timeline, dict):
                document.add_heading("Сроки", level=2)
                document.add_paragraph(f"Оценка длительности: {timeline.get('estimated_duration_days')} дн.")
                document.add_paragraph(f"Confidence: {timeline.get('confidence', 0)}")

        stream = BytesIO()
        document.save(stream)
        stream.seek(0)
        return stream.read()

    def build_cost_workbook(self, project: Project, analysis: Optional[Dict[str, Any]]) -> bytes:
        if not analysis:
            raise ValueError("Analysis data is required to build Excel report.")
        cost = analysis.get("cost")
        if not isinstance(cost, dict):
            raise ValueError("Cost analysis is required to build Excel report.")

        workbook = Workbook()
        summary_ws = workbook.active
        summary_ws.title = "Summary"

        summary_ws["A1"] = "Проект"
        summary_ws["B1"] = f"{project.name} ({project.code})"
        summary_ws["A2"] = "Дата формирования"
        summary_ws["B2"] = datetime.utcnow().strftime("%d.%m.%Y %H:%M")
        summary_ws["A4"] = "Показатель"
        summary_ws["B4"] = "Сумма, руб."
        summary_ws["A4"].font = Font(bold=True)
        summary_ws["B4"].font = Font(bold=True)

        summary_rows = [
            ("Базовая стоимость", cost.get("total_cost", 0)),
            ("Накладные расходы (18%)", cost.get("overhead_cost", 0)),
            ("Прибыль (10%)", cost.get("profit", 0)),
            ("Итого", cost.get("grand_total", cost.get("total_cost", 0))),
            ("Маржа, %", cost.get("margin_percent", 0)),
            ("Итого с командировочными и СИЗ", cost.get("total_with_travel", 0)),
        ]
        row_idx = 5
        for title, value in summary_rows:
            summary_ws.cell(row=row_idx, column=1, value=title)
            summary_ws.cell(row=row_idx, column=2, value=float(value))
            row_idx += 1

        row_idx += 1
        summary_ws.cell(row=row_idx, column=1, value="Структура (материалы / труд / техника)")
        summary_ws.cell(row=row_idx, column=1).font = Font(bold=True)
        row_idx += 1
        for key, value in (cost.get("group_breakdown") or {}).items():
            summary_ws.cell(row=row_idx, column=1, value=key.title())
            summary_ws.cell(row=row_idx, column=2, value=float(value))
            row_idx += 1

        row_idx += 1
        summary_ws.cell(row=row_idx, column=1, value="Стоимость по категориям")
        summary_ws.cell(row=row_idx, column=1).font = Font(bold=True)
        row_idx += 1
        for key, value in (cost.get("by_category") or {}).items():
            summary_ws.cell(row=row_idx, column=1, value=key)
            summary_ws.cell(row=row_idx, column=2, value=float(value))
            row_idx += 1

        entries_ws = workbook.create_sheet(title="Entries")
        headers = [
            "Наименование",
            "Количество",
            "Ед. изм.",
            "Стоимость, руб.",
            "Категория",
            "Группа",
            "Цена за единицу",
            "Источник",
        ]
        entries_ws.append(headers)
        for cell in entries_ws[1]:
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center")

        for entry in cost.get("entries", []):
            entries_ws.append(
                [
                    entry.get("name"),
                    entry.get("quantity"),
                    entry.get("unit"),
                    entry.get("cost"),
                    entry.get("category"),
                    entry.get("group"),
                    entry.get("price_per_unit"),
                    entry.get("source"),
                ]
            )

        labor = analysis.get("labor")
        if isinstance(labor, dict):
            labor_ws = workbook.create_sheet(title="Labor")
            labor_ws.append(["Показатель", "Значение"])
            labor_ws["A1"].font = Font(bold=True)
            labor_ws["B1"].font = Font(bold=True)
            labor_ws.append(["Трудоёмкость, чел·ч", labor.get("total_labor_hours", 0)])
            labor_ws.append(["Эквивалент работников", labor.get("total_worker_equivalent", 0)])
            labor_ws.append(["Эквивалент человеко-дней", labor.get("worker_days", 0)])
            labor_ws.append([])
            labor_ws.append(["График", "Требуемые бригады"])
            labor_ws["A5"].font = Font(bold=True)
            labor_ws["B5"].font = Font(bold=True)
            for schedule, data in (labor.get("schedules") or {}).items():
                labor_ws.append([schedule, data.get("required_workers", 0)])

            labor_ws.append([])
            labor_ws.append(["Детализация"])
            labor_ws["A8"].font = Font(bold=True)
            labor_ws.append(["Работа", "Норма", "Часы", "Работники (экв.)"])
            labor_ws["A9"].font = Font(bold=True)
            labor_ws["B9"].font = Font(bold=True)
            labor_ws["C9"].font = Font(bold=True)
            labor_ws["D9"].font = Font(bold=True)
            for item in labor.get("entries", []):
                labor_ws.append(
                    [
                        item.get("volume_name"),
                        item.get("norm_code"),
                        item.get("labor_hours"),
                        item.get("worker_equivalent"),
                    ]
                )

        stream = BytesIO()
        travel = analysis.get("travel")
        if isinstance(travel, dict):
            travel_ws = workbook.create_sheet(title="Travel & PPE")
            travel_ws.append(["Показатель", "Сумма"])
            travel_ws["A1"].font = Font(bold=True)
            travel_ws["B1"].font = Font(bold=True)
            travel_ws.append(["Командировочные (всего)", travel.get("total_travel", 0)])
            travel_ws.append(["Билеты", travel.get("tickets", 0)])
            travel_ws.append(["Проживание", travel.get("lodging", 0)])
            travel_ws.append(["Суточные", travel.get("per_diem", 0)])
            travel_ws.append(["СИЗ", travel.get("ppe", 0)])
            travel_ws.append(["Итого до коэффициента", travel.get("total_before_coeff", 0)])
            travel_ws.append(["Итого с коэффициентом", travel.get("total_with_coeff", 0)])
            travel_ws.append(["Коэффициент", travel.get("coefficient", 1.0)])
            travel_ws.append(["Работники (чел.)", travel.get("worker_count", 0)])
            travel_ws.append(["Человеко-дни", travel.get("worker_days", 0)])

        financial = analysis.get("financial")
        if isinstance(financial, dict):
            fin_ws = workbook.create_sheet(title="Financial")
            fin_ws.append(["Показатель", "Значение"])
            fin_ws["A1"].font = Font(bold=True)
            fin_ws["B1"].font = Font(bold=True)
            fin_ws.append(["NPV, руб.", financial.get("npv", 0)])
            irr_value = financial.get("irr_percent")
            fin_ws.append(["IRR, %", irr_value if irr_value is not None else "n/a"])
            fin_ws.append(["Срок окупаемости, мес.", financial.get("payback_months", "n/a")])
            assumptions = financial.get("assumptions") or {}
            fin_ws.append([])
            fin_ws.append(["Предположения"])
            fin_ws["A5"].font = Font(bold=True)
            for key, value in assumptions.items():
                fin_ws.append([key, value])
            risk = financial.get("risk") or {}
            fin_ws.append([])
            fin_ws.append(["Риск", "Значение"])
            fin_ws["A" + str(fin_ws.max_row)].font = Font(bold=True)
            fin_ws.append(["Уровень", risk.get("level", "n/a")])
            fin_ws.append(["Score", risk.get("score", 0)])
            fin_ws.append([])
            fin_ws.append(["Факторы риска"])
            fin_ws["A" + str(fin_ws.max_row)].font = Font(bold=True)
            for factor in risk.get("factors", []):
                fin_ws.append([factor])

        workbook.save(stream)
        stream.seek(0)
        return stream.read()

    def generate_preliminary_bundle(self, project: Project, analysis: Optional[Dict[str, Any]]) -> Dict[str, bytes]:
        """
        Возвращает комплект файлов (DOCX, XLSX, PDF) с учётом шаблона Jinja2 и конвертации docx2pdf.
        """
        docx_bytes = self.build_preliminary_report(project, analysis)
        xlsx_bytes = self.build_cost_workbook(project, analysis)

        pdf_bytes: Optional[bytes] = None
        try:
            pdf_bytes = self._convert_docx_to_pdf(docx_bytes)
        except Exception as exc:  # noqa: BLE001
            logger.warning("Failed to convert TEO report to PDF: %s", exc)

        bundle: Dict[str, bytes] = {"docx": docx_bytes, "xlsx": xlsx_bytes}
        if pdf_bytes:
            bundle["pdf"] = pdf_bytes
        return bundle

    def _render_text_summary(self, project: Project, analysis: Dict[str, Any]) -> str:
        try:
            template: Template = self._env.get_template("teo_report_summary.j2")
        except Exception as exc:  # noqa: BLE001
            logger.warning("Unable to load TEO report template: %s", exc)
            return ""

        context = {
            "project": {
                "name": project.name,
                "code": project.code,
                "customer_name": getattr(project, "customer_name", None),
            },
            "analysis": self._normalize_analysis(analysis),
            "generated_at": datetime.utcnow().strftime("%d.%m.%Y %H:%M"),
        }
        return template.render(**context)

    def _normalize_analysis(self, analysis: Dict[str, Any]) -> Dict[str, Any]:
        def ensure_dict(data: Any) -> Dict[str, Any]:
            return data if isinstance(data, dict) else {}

        return {
            "work_volume": ensure_dict(analysis.get("work_volume")),
            "cost": ensure_dict(analysis.get("cost")),
            "labor": ensure_dict(analysis.get("labor")),
            "travel": ensure_dict(analysis.get("travel")),
            "financial": ensure_dict(analysis.get("financial")),
        }

    def _convert_docx_to_pdf(self, docx_bytes: bytes) -> bytes:
        with tempfile.TemporaryDirectory() as tmp_dir:
            tmp_dir_path = Path(tmp_dir)
            docx_path = tmp_dir_path / "report.docx"
            pdf_path = tmp_dir_path / "report.pdf"
            docx_path.write_bytes(docx_bytes)
            convert(str(docx_path), str(pdf_path))
            return pdf_path.read_bytes()

    def _detect_unit(self, work_volume: Dict[str, Any]) -> str:
        entries = work_volume.get("entries") or []
        if entries:
            unit = entries[0].get("unit")
            if unit:
                return unit
        return ""


teo_report_service = TEOReportService()

__all__ = ["teo_report_service", "TEOReportService"]


