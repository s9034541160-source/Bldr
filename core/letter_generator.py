"""
Letter generator module with LM Studio integration
"""

import json
import requests
from typing import Dict, Any, Optional
from datetime import datetime
import os

# Try to import required libraries
try:
    from docx import Document  # type: ignore
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    Document = None

class LetterGenerator:
    def __init__(self, lm_studio_url: str = "http://localhost:1234"):
        self.lm_studio_url = lm_studio_url
        self.lm_studio_available = self._check_lm_studio_connection()
    
    def _check_lm_studio_connection(self) -> bool:
        """Check if LM Studio is available"""
        try:
            response = requests.get(f"{self.lm_studio_url}/v1/models", timeout=5)
            return response.status_code == 200
        except:
            return False
    
    def generate_letter_with_lm(self, 
                               description: str, 
                               template_text: str, 
                               project_data: Optional[Dict[str, Any]] = None,
                               tone: float = 0.0,
                               dryness: float = 0.5,
                               humanity: float = 0.7,
                               length: str = "medium",
                               formality: str = "formal") -> str:
        """
        Generate a letter using LM Studio
        
        Args:
            description: Description of the issue/problem
            template_text: Template text to use as base
            project_data: Optional project data to include
            tone: Tone parameter (-1 to +1: harsh to loyal)
            dryness: Dryness parameter (0-1: lively to dry)
            humanity: Humanity parameter (0-1: robotic to natural)
            length: Length parameter (short/medium/long)
            formality: Formality parameter (formal/informal)
            
        Returns:
            Generated letter text
        """
        if not self.lm_studio_available:
            return self._generate_fallback_letter(description, template_text, project_data)
        
        # Prepare the prompt
        system_prompt = f"""You are a professional RU letter writer for construction company Bldr Empire. 
Generate formal business letter in Russian. 

Parameters:
- Tone {tone} (-1 negative/harsh, 0 neutral, +1 loyal/benevolent)
- Dryness {dryness} (0 lively, 1 dry/formal)
- Humanity {humanity} (0 robotic, 1 natural/human)
- Length {length} (short<300w, medium 500w, long>800w)
- Formality {formality} (formal/informal)

Use placeholders: [Получатель], [Дата], [Отправитель].
Structure: Header (From/To/Subject/Date), Body (Intro/Problem/Solution/Call to action), Footer (Signature).

Based on description: {description}
"""
        
        # Prepare the user prompt with template and project data
        user_prompt = template_text
        
        # Add project data if available
        if project_data:
            user_prompt += f"\n\nFrom project {project_data.get('name', '')}: "
            if project_data.get('violations'):
                user_prompt += f"violations {', '.join(project_data.get('violations', []))}"
        
        # Prepare the messages
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        
        # Map length to max tokens
        length_tokens = {
            "short": 1000,
            "medium": 2000,
            "long": 4000
        }
        
        # Adjust temperature based on humanity parameter
        temperature = 0.4 + (humanity * 0.4)  # Range 0.4-0.8
        
        # Prepare the request payload
        payload = {
            "model": os.getenv("DEFAULT_LM_MODEL", "llama3"),  # Default model, can be changed
            "messages": messages,
            "temperature": temperature,
            "max_tokens": length_tokens.get(length, 1000),
            "stream": False
        }
        
        try:
            # Make the API call to LM Studio
            response = requests.post(
                f"{self.lm_studio_url}/v1/chat/completions",
                json=payload,
                timeout=30
            )
            
            if response.status_code == 200:
                result = response.json()
                choices = result.get('choices') if isinstance(result, dict) else None
                if isinstance(choices, list) and choices:
                    first = choices[0] if isinstance(choices[0], dict) else {}
                    msg = first.get('message') if isinstance(first, dict) else None
                    if isinstance(msg, dict):
                        generated_text = msg.get('content') or ''
                    else:
                        generated_text = first.get('text') if isinstance(first, dict) else ''
                else:
                    generated_text = result.get('content') or result.get('message') or ''
                return generated_text
            else:
                # Fallback to template-based generation
                return self._generate_fallback_letter(description, template_text, project_data)
                
        except Exception as e:
            print(f"Error calling LM Studio: {e}")
            # Fallback to template-based generation
            return self._generate_fallback_letter(description, template_text, project_data)
    
    def improve_letter_draft(self,
                            draft: str,
                            description: str = "",
                            template_text: str = "",
                            project_data: Optional[Dict[str, Any]] = None,
                            tone: float = 0.0,
                            dryness: float = 0.5,
                            humanity: float = 0.7,
                            length: str = "medium",
                            formality: str = "formal") -> str:
        """
        Improve an existing letter draft using LM Studio
        
        Args:
            draft: Existing letter draft to improve
            description: Description of the issue/problem
            template_text: Template text to use as reference
            project_data: Optional project data to include
            tone: Tone parameter (-1 to +1: harsh to loyal)
            dryness: Dryness parameter (0-1: lively to dry)
            humanity: Humanity parameter (0-1: robotic to natural)
            length: Length parameter (short/medium/long)
            formality: Formality parameter (formal/informal)
            
        Returns:
            Improved letter text
        """
        if not self.lm_studio_available:
            return self._improve_fallback_letter(draft, description, project_data)
        
        # Prepare the prompt
        system_prompt = f"""Improve this draft letter for construction company. 
Apply parameters: tone {tone}, dryness {dryness}, humanity {humanity}, length {length}, formality {formality}.
Make it professional RU business letter. Fix grammar, add structure, incorporate description: {description}.
Use placeholders: [Получатель], [Дата], [Отправитель].
Structure: Header (From/To/Subject/Date), Body (Intro/Problem/Solution/Call to action), Footer (Signature).
Language: Russian only.
"""
        
        # Prepare the user prompt with draft
        user_prompt = f"Draft to improve:\n{draft}"
        
        # Add project data if available
        if project_data:
            user_prompt += f"\n\nFrom project {project_data.get('name', '')}: "
            if project_data.get('violations'):
                user_prompt += f"violations {', '.join(project_data.get('violations', []))}"
        
        # Prepare the messages
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ]
        
        # Map length to max tokens
        length_tokens = {
            "short": 1000,
            "medium": 2000,
            "long": 4000
        }
        
        # Adjust temperature based on humanity parameter
        temperature = 0.4 + (humanity * 0.4)  # Range 0.4-0.8
        
        # Prepare the request payload
        payload = {
            "model": os.getenv("DEFAULT_LM_MODEL", "llama3"),  # Default model, can be changed
            "messages": messages,
            "temperature": temperature,
            "max_tokens": length_tokens.get(length, 1000),
            "stream": False
        }
        
        try:
            # Make the API call to LM Studio
            response = requests.post(
                f"{self.lm_studio_url}/v1/chat/completions",
                json=payload,
                timeout=30
            )
            
            if response.status_code == 200:
                result = response.json()
                choices = result.get('choices') if isinstance(result, dict) else None
                if isinstance(choices, list) and choices:
                    first = choices[0] if isinstance(choices[0], dict) else {}
                    msg = first.get('message') if isinstance(first, dict) else None
                    if isinstance(msg, dict):
                        improved_text = msg.get('content') or ''
                    else:
                        improved_text = first.get('text') if isinstance(first, dict) else ''
                else:
                    improved_text = result.get('content') or result.get('message') or ''
                return improved_text
            else:
                # Fallback to template-based improvement
                return self._improve_fallback_letter(draft, description, project_data)
                
        except Exception as e:
            print(f"Error calling LM Studio: {e}")
            # Fallback to template-based improvement
            return self._improve_fallback_letter(draft, description, project_data)
    
    def _generate_fallback_letter(self, 
                                 description: str, 
                                 template_text: str, 
                                 project_data: Optional[Dict[str, Any]] = None) -> str:
        """
        Generate a fallback letter when LM Studio is not available
        """
        # Simple template-based generation
        letter = template_text.replace("[Получатель]", project_data.get("client", "Заказчик") if project_data else "Заказчик")
        letter = letter.replace("[Дата]", datetime.now().strftime("%d.%m.%Y"))
        letter = letter.replace("[Отправитель]", "АО БЛДР")
        
        # Add description if not already in template
        if description and description not in letter:
            letter += f"\n\nОписание проблемы: {description}"
        
        # Add project data if available
        if project_data:
            letter += f"\n\nПроект: {project_data.get('name', 'Не указан')}"
            if project_data.get('violations'):
                letter += f"\nНарушения: {', '.join(project_data.get('violations', []))}"
        
        return letter
    
    def _improve_fallback_letter(self,
                                draft: str,
                                description: str = "",
                                project_data: Optional[Dict[str, Any]] = None) -> str:
        """
        Improve a letter draft using fallback method when LM Studio is not available
        """
        # Simple improvement - just add project data if not already present
        improved = draft
        
        if project_data:
            if project_data.get('name') and project_data['name'] not in improved:
                improved += f"\n\nПроект: {project_data['name']}"
            if project_data.get('violations') and not any(v in improved for v in project_data['violations']):
                improved += f"\nНарушения: {', '.join(project_data['violations'])}"
        
        if description and description not in improved:
            improved += f"\n\nОписание проблемы: {description}"
        
        return improved
    
    def export_to_docx(self, letter_text: str, filename: Optional[str] = None) -> str:
        """
        Export letter to DOCX format
        
        Args:
            letter_text: Letter text to export
            filename: Optional filename (will generate if not provided)
            
        Returns:
            Path to the generated DOCX file
        """
        if not HAS_PYTHON_DOCX:
            raise ImportError("python-docx library is required for DOCX export")
        
        if not filename:
            filename = f"letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # Create document
        if Document is not None:
            doc = Document()
        else:
            raise ImportError("python-docx library is required for DOCX export")
        
        # Add title
        doc.add_heading('Официальное письмо', 0)
        
        # Split text into paragraphs and add to document
        paragraphs = letter_text.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                # Check if it's a heading
                if paragraph.startswith('#') or len(paragraph) < 100:
                    doc.add_heading(paragraph.replace('#', '').strip(), level=1)
                else:
                    doc.add_paragraph(paragraph)
        
        # Save document
        doc.save(filename)
        return filename

# Initialize letter generator
letter_generator = LetterGenerator()