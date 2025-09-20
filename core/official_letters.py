"""
Official letters generator module with real Jinja2 templates and docx2pdf export
"""

import json
from typing import Dict, Any
from datetime import datetime
import os

# Try to import required libraries
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False
    Document = None
    Inches = None
    WD_ALIGN_PARAGRAPH = None

# Try to import Jinja2
try:
    from jinja2 import Environment, FileSystemLoader
    HAS_JINJA2 = True
except ImportError:
    HAS_JINJA2 = False
    Environment = None
    FileSystemLoader = None

# Try to import docx2pdf
try:
    from docx2pdf import convert
    HAS_DOCX2PDF = True
except ImportError:
    HAS_DOCX2PDF = False
    convert = None

def generate_official_letter(template: str, data: Dict[str, Any]) -> str:
    """
    Generate an official letter using real Jinja2 templates and docx2pdf export
    
    Args:
        template: Template name or type of letter
        data: Data to populate the letter
        
    Returns:
        Path to the generated letter file
    """
    # Check if required libraries are available
    if not HAS_PYTHON_DOCX:
        raise ImportError("python-docx library is required for letter generation")
    
    # Use real Jinja2 templates if available
    if HAS_JINJA2 and os.path.exists('templates'):
        try:
            # Create Jinja2 environment
            env = Environment(loader=FileSystemLoader('templates'))
            
            # Load template
            template_file = f"{template}.docx.j2"
            if not os.path.exists(f"templates/{template_file}"):
                template_file = "base_letter.docx.j2"
            
            jinja_template = env.get_template(template_file)
            
            # Add default data
            template_data = {
                "date": datetime.now().strftime("%d.%m.%Y"),
                "recipient": data.get("recipient", ""),
                "sender": data.get("sender", ""),
                "subject": data.get("subject", ""),
                "signer": data.get("signer", ""),
                "total_cost": data.get("total_cost", ""),
                "compliance_details": data.get("compliance_details", []),
                "violations": data.get("violations", []),
                "budget_details": data.get("budget_details", {}),
                "content": data.get("content", "")
            }
            
            # Render template
            rendered_content = jinja_template.render(**template_data)
            
            # Create document from template
            doc = Document()
            
            # Parse rendered content and add to document
            # For simplicity, we'll add it as paragraphs
            for line in rendered_content.split('\n'):
                if line.strip():
                    if line.startswith('<h1>'):
                        doc.add_heading(line.replace('<h1>', '').replace('</h1>', ''), level=1)
                    elif line.startswith('<h2>'):
                        doc.add_heading(line.replace('<h2>', '').replace('</h2>', ''), level=2)
                    elif line.startswith('<p>'):
                        doc.add_paragraph(line.replace('<p>', '').replace('</p>', ''))
                    elif line.startswith('<ul>'):
                        continue  # Skip ul tags
                    elif line.startswith('<li>'):
                        doc.add_paragraph(line.replace('<li>', '').replace('</li>', ''), style='List Bullet')
                    else:
                        doc.add_paragraph(line)
        except Exception as e:
            print(f"Error using Jinja2 template: {e}")
            # Fallback to simple docx generation
            doc = Document()
    else:
        # Fallback to simple docx generation
        doc = Document()
    
    # Add header
    doc.add_heading('Официальное письмо', 0)
    
    # Add date
    today = datetime.now().strftime("%d.%m.%Y")
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run(f"Дата: {today}")
    
    # Add recipient
    if 'recipient' in data:
        doc.add_paragraph(f"Кому: {data['recipient']}")
    
    # Add sender
    if 'sender' in data:
        doc.add_paragraph(f"От: {data['sender']}")
    
    # Add subject
    if 'subject' in data:
        doc.add_paragraph(f"Тема: {data['subject']}")
    
    # Add content based on template
    if template == 'compliance_sp31':
        doc.add_heading('Письмо о соответствии СП31', level=1)
        doc.add_paragraph(
            "Настоящим письмом подтверждаем соответствие проектной документации "
            "требованиям СП 45.13330.2017 (актуализированная редакция СНиП 3.02.01-87*) "
            "в части организации строительного контроля."
        )
        
        # Add compliance details
        if 'compliance_details' in data:
            doc.add_heading('Подробности соответствия:', level=2)
            for detail in data['compliance_details']:
                doc.add_paragraph(detail, style='List Bullet')
        
        # Add violations if any
        if 'violations' in data and data['violations']:
            doc.add_heading('Выявленные нарушения:', level=2)
            for violation in data['violations']:
                doc.add_paragraph(violation, style='List Bullet')
        else:
            doc.add_paragraph("Нарушений не выявлено.")
            
    elif template == 'budget_approval':
        doc.add_heading('Письмо о согласовании сметы', level=1)
        doc.add_paragraph(
            "Настоящим письмом направляем на согласование сметную документацию "
            "по объекту строительства."
        )
        
        # Add budget details
        if 'budget_details' in data:
            doc.add_heading('Сметные показатели:', level=2)
            doc.add_paragraph(f"Общая стоимость: {data['budget_details'].get('total_cost', 'Не указана')}")
            doc.add_paragraph(f"Сроки выполнения: {data['budget_details'].get('duration', 'Не указаны')}")
            
    else:
        # Generic letter
        doc.add_heading('Основное содержание', level=1)
        if 'content' in data:
            doc.add_paragraph(data['content'])
    
    # Add signature section
    doc.add_paragraph("\n")
    if 'signer' in data:
        doc.add_paragraph(f"Ответственное лицо: {data['signer']}")
    doc.add_paragraph("Подпись: ________________")
    
    # Save document
    filename = f"official_letter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    
    # Convert to PDF if docx2pdf is available
    pdf_filename = filename.replace('.docx', '.pdf')
    if HAS_DOCX2PDF and convert is not None:
        try:
            convert(filename, pdf_filename)
            return pdf_filename  # Return PDF file path
        except Exception as e:
            print(f"Error converting to PDF: {e}")
    
    return filename

def get_letter_templates() -> Dict[str, str]:
    """
    Get available letter templates with real template discovery
    
    Returns:
        Dictionary of template names and descriptions
    """
    templates = {
        "compliance_sp31": "Письмо о соответствии СП31",
        "budget_approval": "Письмо о согласовании сметы",
        "work_completion": "Письмо о выполнении работ",
        "technical_spec": "Письмо с техническими спецификациями",
        "contract_proposal": "Договорное предложение"
    }
    
    # Discover real templates from templates directory
    if HAS_JINJA2 and os.path.exists('templates'):
        try:
            template_files = os.listdir('templates')
            for template_file in template_files:
                if template_file.endswith('.j2') or template_file.endswith('.docx.j2'):
                    template_name = template_file.replace('.docx.j2', '').replace('.j2', '')
                    if template_name not in templates:
                        templates[template_name] = f"Пользовательский шаблон: {template_name}"
        except Exception as e:
            print(f"Error discovering templates: {e}")
    
    return templates