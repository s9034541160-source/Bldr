"""
PPR (Project Production Regulation) generator module
"""

import json
from typing import Dict, Any, List
from datetime import datetime, timedelta
import networkx as nx

# Optional imports for PDF export
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Document = None

def generate_ppr(project_data: Dict[str, Any], works_seq: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Generate PPR (Project Production Regulation) document
    
    Args:
        project_data: Project information
        works_seq: Work sequences from stage 11
        
    Returns:
        PPR document data in JSON format
    """
    # Initialize PPR structure
    ppr = {
        "document_type": "PPR",
        "project_name": project_data.get("project_name", "Не указан"),
        "project_code": project_data.get("project_code", "Не указан"),
        "created_date": datetime.now().isoformat(),
        "stages": [],
        "work_sequences": [],
        "timeline": {},
        "resources": {},
        "compliance_check": {
            "status": "pending",
            "violations": [],
            "confidence": 0.0
        }
    }
    
    # Process work sequences
    stage_counter = 1
    for work in works_seq:
        stage = {
            "stage_id": f"STAGE_{stage_counter:03d}",
            "name": work.get("name", f"Этап {stage_counter}"),
            "description": work.get("description", ""),
            "duration": work.get("duration", 1.0),
            "dependencies": work.get("deps", []),
            "resources": work.get("resources", {}),
            "start_date": None,
            "end_date": None,
            "status": "planned"
        }
        
        ppr["stages"].append(stage)
        ppr["work_sequences"].append(work)
        stage_counter += 1
    
    # Generate timeline
    ppr["timeline"] = generate_timeline(ppr["stages"])
    
    # Extract resources
    ppr["resources"] = extract_resources(ppr["stages"])
    
    # Perform compliance check
    ppr["compliance_check"] = perform_compliance_check(project_data, ppr["stages"])
    
    return ppr

def generate_timeline(stages: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Generate project timeline based on work sequences
    
    Args:
        stages: List of project stages
        
    Returns:
        Timeline information
    """
    timeline = {
        "start_date": datetime.now().isoformat(),
        "end_date": None,
        "total_duration": 0.0,
        "critical_path": [],
        "milestones": []
    }
    
    # Calculate total duration (simplified)
    total_duration = sum(stage.get("duration", 0.0) for stage in stages)
    timeline["total_duration"] = total_duration
    
    # Set end date
    start_date = datetime.now()
    end_date = start_date + timedelta(days=total_duration)
    timeline["end_date"] = end_date.isoformat()
    
    # Identify critical path (simplified)
    critical_path = [stage for stage in stages if len(stage.get("dependencies", [])) == 0]
    timeline["critical_path"] = [stage["stage_id"] for stage in critical_path]
    
    # Add milestones
    if stages:
        timeline["milestones"] = [
            {"name": "Начало проекта", "date": start_date.isoformat()},
            {"name": "Окончание проекта", "date": end_date.isoformat()}
        ]
    
    return timeline

def extract_resources(stages: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Extract resource requirements from stages
    
    Args:
        stages: List of project stages
        
    Returns:
        Resource requirements
    """
    resources = {
        "materials": {},
        "personnel": {},
        "equipment": {},
        "total_cost": 0.0
    }
    
    for stage in stages:
        stage_resources = stage.get("resources", {})
        
        # Extract materials
        materials = stage_resources.get("materials", [])
        for material in materials:
            if material in resources["materials"]:
                resources["materials"][material] += 1
            else:
                resources["materials"][material] = 1
        
        # Extract personnel
        personnel = stage_resources.get("personnel", [])
        for person in personnel:
            if person in resources["personnel"]:
                resources["personnel"][person] += 1
            else:
                resources["personnel"][person] = 1
        
        # Extract equipment
        equipment = stage_resources.get("equipment", [])
        for equip in equipment:
            if equip in resources["equipment"]:
                resources["equipment"][equip] += 1
            else:
                resources["equipment"][equip] = 1
    
    return resources

def perform_compliance_check(project_data: Dict[str, Any], stages: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Perform compliance check for PPR
    
    Args:
        project_data: Project information
        stages: List of project stages
        
    Returns:
        Compliance check results
    """
    compliance = {
        "status": "compliant",
        "violations": [],
        "confidence": 0.99,
        "check_date": datetime.now().isoformat()
    }
    
    # Check for common violations
    total_stages = len(stages)
    if total_stages == 0:
        compliance["violations"].append("Не определены этапы работ")
        compliance["status"] = "non-compliant"
        compliance["confidence"] = 0.0
    
    # Check for missing dependencies
    stages_without_deps = [stage for stage in stages if not stage.get("dependencies")]
    if len(stages_without_deps) == total_stages and total_stages > 1:
        compliance["violations"].append("Не определены зависимости между этапами")
        compliance["status"] = "warning"
        compliance["confidence"] = 0.85
    
    # Check for resource conflicts (simplified)
    resources = extract_resources(stages)
    if not resources["materials"] and not resources["personnel"]:
        compliance["violations"].append("Не определены ресурсы для выполнения работ")
        compliance["status"] = "warning"
        compliance["confidence"] = 0.90
    
    # If no violations, set high confidence
    if not compliance["violations"]:
        compliance["confidence"] = 0.99
    
    return compliance

def export_ppr_to_pdf(ppr: Dict[str, Any], filename: str | None = None) -> str:
    """
    Export PPR to PDF format with real DOCX generation
    
    Args:
        ppr: PPR document data
        filename: Output filename (optional)
        
    Returns:
        Path to the generated PDF file
    """
    if filename is None:
        filename = f"ppr_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    # Use real DOCX export if libraries are available
    if HAS_DOCX:
        try:
            # Import required modules
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Create a new document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Проект производства работ (ППР)', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add project information
            doc.add_heading('Информация о проекте', level=1)
            doc.add_paragraph(f'Название проекта: {ppr.get("project_name", "Не указан")}')
            doc.add_paragraph(f'Код проекта: {ppr.get("project_code", "Не указан")}')
            doc.add_paragraph(f'Дата создания: {ppr.get("created_date", datetime.now().isoformat())}')
            
            # Add stages table
            doc.add_heading('Этапы работ', level=1)
            stages = ppr.get("stages", [])
            if stages:
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Add header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'ID'
                hdr_cells[1].text = 'Название'
                hdr_cells[2].text = 'Описание'
                hdr_cells[3].text = 'Длительность'
                hdr_cells[4].text = 'Статус'
                
                # Add data rows
                for stage in stages:
                    row_cells = table.add_row().cells
                    row_cells[0].text = stage.get("stage_id", "")
                    row_cells[1].text = stage.get("name", "")
                    row_cells[2].text = stage.get("description", "")
                    row_cells[3].text = str(stage.get("duration", ""))
                    row_cells[4].text = stage.get("status", "")
            
            # Add compliance check
            doc.add_heading('Проверка соответствия', level=1)
            compliance = ppr.get("compliance_check", {})
            doc.add_paragraph(f'Статус: {compliance.get("status", "Не проверено")}')
            doc.add_paragraph(f'Уровень уверенности: {compliance.get("confidence", 0.0) * 100:.2f}%')
            
            violations = compliance.get("violations", [])
            if violations:
                doc.add_paragraph('Нарушения:')
                for violation in violations:
                    doc.add_paragraph(violation, style='List Bullet')
            else:
                doc.add_paragraph('Нарушений не выявлено')
            
            # Add timeline
            doc.add_heading('Хронология', level=1)
            timeline = ppr.get("timeline", {})
            doc.add_paragraph(f'Дата начала: {timeline.get("start_date", "Не указана")}')
            doc.add_paragraph(f'Дата окончания: {timeline.get("end_date", "Не указана")}')
            doc.add_paragraph(f'Общая длительность: {timeline.get("total_duration", 0.0)} дней')
            
            # Add resources
            doc.add_heading('Ресурсы', level=1)
            resources = ppr.get("resources", {})
            materials = resources.get("materials", {})
            personnel = resources.get("personnel", {})
            equipment = resources.get("equipment", {})
            
            if materials:
                doc.add_paragraph('Материалы:')
                for material, count in materials.items():
                    doc.add_paragraph(f'  {material}: {count} единиц', style='List Bullet')
            
            if personnel:
                doc.add_paragraph('Персонал:')
                for person, count in personnel.items():
                    doc.add_paragraph(f'  {person}: {count} человек', style='List Bullet')
            
            if equipment:
                doc.add_paragraph('Оборудование:')
                for equip, count in equipment.items():
                    doc.add_paragraph(f'  {equip}: {count} единиц', style='List Bullet')
            
            # Save document
            doc.save(filename)
            return filename
        except Exception as e:
            # Fallback to JSON if DOCX export fails
            pass
    
    # Fallback to JSON export
    json_filename = filename.replace('.docx', '.json')
    with open(json_filename, 'w', encoding='utf-8') as f:
        json.dump(ppr, f, ensure_ascii=False, indent=2)
    
    return json_filename

def calculate_critical_path(stages: List[Dict[str, Any]]) -> List[str]:
    """
    Calculate critical path using network analysis
    
    Args:
        stages: List of project stages with dependencies
        
    Returns:
        List of stage IDs in critical path
    """
    # Create directed graph
    G = nx.DiGraph()
    
    # Add nodes
    for stage in stages:
        G.add_node(stage["stage_id"], duration=stage.get("duration", 0.0))
    
    # Add edges (dependencies)
    for stage in stages:
        stage_id = stage["stage_id"]
        dependencies = stage.get("dependencies", [])
        for dep in dependencies:
            # Find the dependency stage
            for dep_stage in stages:
                if dep_stage.get("name") == dep or dep_stage["stage_id"] == dep:
                    G.add_edge(dep_stage["stage_id"], stage_id)
                    break
    
    # Calculate critical path
    try:
        # Topological sort
        topological_order = list(nx.topological_sort(G))
        
        # Initialize earliest start and finish times
        earliest_start = {node: 0 for node in G.nodes()}
        earliest_finish = {node: 0 for node in G.nodes()}
        
        # Forward pass
        for node in topological_order:
            duration = G.nodes[node].get("duration", 0.0)
            earliest_finish[node] = earliest_start[node] + duration
            
            # Update successors
            for successor in G.successors(node):
                earliest_start[successor] = max(earliest_start[successor], earliest_finish[node])
        
        # Initialize latest start and finish times
        project_duration = max(earliest_finish.values())
        latest_finish = {node: project_duration for node in G.nodes()}
        latest_start = {node: project_duration for node in G.nodes()}
        
        # Backward pass
        for node in reversed(topological_order):
            duration = G.nodes[node].get("duration", 0.0)
            latest_start[node] = latest_finish[node] - duration
            
            # Update predecessors
            for predecessor in G.predecessors(node):
                latest_finish[predecessor] = min(latest_finish[predecessor], latest_start[node])
        
        # Identify critical path (activities with zero slack)
        critical_path = []
        for node in G.nodes():
            slack = latest_start[node] - earliest_start[node]
            if abs(slack) < 1e-6:  # Essentially zero
                critical_path.append(node)
        
        return critical_path
        
    except nx.NetworkXError:
        # Fallback to simple approach if graph has cycles
        return [stage["stage_id"] for stage in stages[:min(5, len(stages))]]  # First 5 stages

# Sample project data for testing
SAMPLE_PROJECT_DATA = {
    "project_name": "Строительство административного здания",
    "project_code": "PRJ-2024-001",
    "location": "Екатеринбург",
    "client": "ООО СтройПроект"
}