#!/usr/bin/env python3
"""
Advanced Template Analysis System
Система анализа шаблонов документов с автоматической категоризацией и извлечением метаданных
"""

import os
import re
import json
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
import hashlib
import logging

# Импорты для обработки различных форматов документов
try:
    import docx
    from docx.document import Document as DocxDocument
except ImportError:
    docx = None

try:
    import PyPDF2
    import pdfplumber
except ImportError:
    PyPDF2 = None
    pdfplumber = None

try:
    from odf import text, teletype
    from odf.opendocument import load as odf_load
except ImportError:
    text = None
    odf_load = None

logger = logging.getLogger(__name__)

class DocumentFormat(Enum):
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    RTF = "rtf"
    ODT = "odt"
    TXT = "txt"
    HTML = "html"
    UNKNOWN = "unknown"

class DocumentCategory(Enum):
    CONTRACT = "contract"
    REPORT = "report"
    APPLICATION = "application"
    LETTER = "letter"
    ACT = "act"
    ESTIMATE = "estimate"
    FORM = "form"
    TEMPLATE = "template"
    LEGAL = "legal"
    TECHNICAL = "technical"
    FINANCIAL = "financial"
    ADMINISTRATIVE = "administrative"

class ComplexityLevel(Enum):
    SIMPLE = "simple"
    MEDIUM = "medium"
    COMPLEX = "complex"
    VERY_COMPLEX = "very_complex"

@dataclass
class PlaceholderInfo:
    name: str
    type: str  # text, number, date, email, phone, etc.
    pattern: str
    required: bool = True
    description: str = ""
    default_value: Optional[str] = None
    validation_rules: List[str] = field(default_factory=list)

@dataclass
class TemplateStructure:
    sections: List[str] = field(default_factory=list)
    headers: List[str] = field(default_factory=list)
    tables: int = 0
    lists: int = 0
    images: int = 0
    pages: int = 1
    paragraphs: int = 0

@dataclass
class TemplateAnalysisResult:
    file_path: str
    format: DocumentFormat
    category: DocumentCategory
    complexity: ComplexityLevel
    
    # Основная информация
    title: str = ""
    description: str = ""
    language: str = "unknown"
    
    # Структурная информация
    structure: TemplateStructure = field(default_factory=TemplateStructure)
    
    # Плейсхолдеры и поля
    placeholders: List[PlaceholderInfo] = field(default_factory=list)
    estimated_fields: int = 0
    
    # Метаданные
    file_size: int = 0
    creation_date: Optional[datetime] = None
    modification_date: Optional[datetime] = None
    author: str = ""
    organization: str = ""
    
    # Качественные характеристики
    quality_score: float = 0.0
    completeness_score: float = 0.0
    usability_score: float = 0.0
    
    # Теги и ключевые слова
    tags: List[str] = field(default_factory=list)
    keywords: List[str] = field(default_factory=list)
    
    # Техническая информация
    analysis_date: datetime = field(default_factory=datetime.now)
    processing_time: float = 0.0
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)class Adva
ncedTemplateAnalyzer:
    """Расширенный анализатор шаблонов документов"""
    
    def __init__(self):
        # Паттерны для определения категорий документов
        self.category_patterns = {
            DocumentCategory.CONTRACT: [
                r'договор', r'контракт', r'соглашение', r'сделка',
                r'contract', r'agreement', r'deal'
            ],
            DocumentCategory.REPORT: [
                r'отчет', r'справка', r'заключение', r'доклад',
                r'report', r'summary', r'conclusion'
            ],
            DocumentCategory.APPLICATION: [
                r'заявление', r'заявка', r'ходатайство', r'прошение',
                r'application', r'request', r'petition'
            ],
            DocumentCategory.LETTER: [
                r'письмо', r'уведомление', r'извещение',
                r'letter', r'notification', r'notice'
            ],
            DocumentCategory.ACT: [
                r'акт', r'протокол', r'свидетельство',
                r'act', r'protocol', r'certificate'
            ],
            DocumentCategory.ESTIMATE: [
                r'смета', r'калькуляция', r'расчет', r'оценка',
                r'estimate', r'calculation', r'assessment'
            ],
            DocumentCategory.FORM: [
                r'форма', r'бланк', r'анкета',
                r'form', r'blank', r'questionnaire'
            ],
            DocumentCategory.TEMPLATE: [
                r'шаблон', r'образец', r'макет',
                r'template', r'sample', r'pattern'
            ]
        }
        
        # Паттерны для поиска плейсхолдеров
        self.placeholder_patterns = [
            # Фигурные скобки
            (r'\{([^}]+)\}', 'curly_braces'),
            # Квадратные скобки
            (r'\[([^\]]+)\]', 'square_brackets'),
            # Подчеркивания
            (r'_{3,}', 'underscores'),
            # Точки
            (r'\.{3,}', 'dots'),
            # Специальные маркеры
            (r'<<([^>]+)>>', 'angle_brackets'),
            (r'%([^%]+)%', 'percent_signs'),
            # Типичные поля
            (r'(ФИО|Ф\.И\.О\.)', 'name_field'),
            (r'(дата|число)', 'date_field'),
            (r'(подпись)', 'signature_field')
        ]
        
        # Ключевые слова для определения языка
        self.language_keywords = {
            'ru': ['и', 'в', 'на', 'с', 'по', 'для', 'от', 'до', 'при', 'что', 'как', 'где'],
            'en': ['the', 'and', 'of', 'to', 'in', 'for', 'with', 'on', 'at', 'by', 'from']
        }
        
        # Паттерны для извлечения структурных элементов
        self.structure_patterns = {
            'headers': [
                r'^[А-ЯA-Z][^.!?]*$',  # Заголовки (заглавные буквы)
                r'^\d+\.\s*[А-ЯA-Z]',   # Нумерованные заголовки
                r'^[IVX]+\.\s*[А-ЯA-Z]' # Римские цифры
            ],
            'sections': [
                r'раздел\s+\d+',
                r'глава\s+\d+',
                r'пункт\s+\d+',
                r'section\s+\d+',
                r'chapter\s+\d+'
            ]
        }
    
    def analyze_template(self, file_path: str) -> TemplateAnalysisResult:
        """Полный анализ шаблона документа"""
        start_time = datetime.now()
        
        try:
            file_path_obj = Path(file_path)
            
            if not file_path_obj.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Определяем формат файла
            document_format = self._detect_format(file_path_obj)
            
            # Создаем базовый результат анализа
            result = TemplateAnalysisResult(
                file_path=str(file_path_obj),
                format=document_format,
                category=DocumentCategory.TEMPLATE,  # По умолчанию
                complexity=ComplexityLevel.MEDIUM,   # По умолчанию
                file_size=file_path_obj.stat().st_size,
                modification_date=datetime.fromtimestamp(file_path_obj.stat().st_mtime)
            )
            
            # Извлекаем текстовое содержимое
            text_content = self._extract_text_content(file_path_obj, document_format)
            
            if text_content:
                # Анализируем содержимое
                self._analyze_content(result, text_content)
                
                # Определяем категорию документа
                result.category = self._categorize_document(text_content)
                
                # Анализируем структуру
                result.structure = self._analyze_structure(text_content)
                
                # Ищем плейсхолдеры
                result.placeholders = self._extract_placeholders(text_content)
                result.estimated_fields = len(result.placeholders)
                
                # Определяем язык
                result.language = self._detect_language(text_content)
                
                # Извлекаем ключевые слова и теги
                result.keywords = self._extract_keywords(text_content)
                result.tags = self._generate_tags(result)
                
                # Оцениваем сложность
                result.complexity = self._assess_complexity(result)
                
                # Вычисляем оценки качества
                self._calculate_quality_scores(result, text_content)
                
            else:
                result.errors.append("Could not extract text content from document")
            
            # Извлекаем метаданные файла
            self._extract_file_metadata(result, file_path_obj, document_format)
            
        except Exception as e:
            logger.error(f"Error analyzing template {file_path}: {e}")
            result.errors.append(str(e))
        
        # Записываем время обработки
        result.processing_time = (datetime.now() - start_time).total_seconds()
        
        return result
    
    def _detect_format(self, file_path: Path) -> DocumentFormat:
        """Определение формата документа"""
        suffix = file_path.suffix.lower()
        
        format_mapping = {
            '.pdf': DocumentFormat.PDF,
            '.docx': DocumentFormat.DOCX,
            '.doc': DocumentFormat.DOC,
            '.rtf': DocumentFormat.RTF,
            '.odt': DocumentFormat.ODT,
            '.txt': DocumentFormat.TXT,
            '.html': DocumentFormat.HTML,
            '.htm': DocumentFormat.HTML
        }
        
        return format_mapping.get(suffix, DocumentFormat.UNKNOWN)
    
    def _extract_text_content(self, file_path: Path, document_format: DocumentFormat) -> str:
        """Извлечение текстового содержимого из документа"""
        try:
            if document_format == DocumentFormat.PDF:
                return self._extract_pdf_text(file_path)
            elif document_format == DocumentFormat.DOCX:
                return self._extract_docx_text(file_path)
            elif document_format == DocumentFormat.DOC:
                return self._extract_doc_text(file_path)
            elif document_format == DocumentFormat.RTF:
                return self._extract_rtf_text(file_path)
            elif document_format == DocumentFormat.ODT:
                return self._extract_odt_text(file_path)
            elif document_format == DocumentFormat.TXT:
                return self._extract_txt_text(file_path)
            elif document_format == DocumentFormat.HTML:
                return self._extract_html_text(file_path)
            else:
                logger.warning(f"Unsupported format: {document_format}")
                return ""
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return "" 
   
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Извлечение текста из PDF"""
        text = ""
        
        # Пробуем pdfplumber (лучше для таблиц и структуры)
        if pdfplumber:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                return text
            except Exception as e:
                logger.warning(f"pdfplumber failed for {file_path}: {e}")
        
        # Fallback на PyPDF2
        if PyPDF2:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                return text
            except Exception as e:
                logger.warning(f"PyPDF2 failed for {file_path}: {e}")
        
        return ""
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Извлечение текста из DOCX"""
        if not docx:
            logger.warning("python-docx not available")
            return ""
        
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""
    
    def _extract_doc_text(self, file_path: Path) -> str:
        """Извлечение текста из DOC (старый формат Word)"""
        # Для .doc файлов нужны специальные библиотеки
        # Пока возвращаем пустую строку
        logger.warning("DOC format extraction not implemented")
        return ""
    
    def _extract_rtf_text(self, file_path: Path) -> str:
        """Извлечение текста из RTF"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                # Простое удаление RTF тегов (упрощенная версия)
                text = re.sub(r'\\[a-z]+\d*\s?', '', content)
                text = re.sub(r'[{}]', '', text)
                return text
        except Exception as e:
            logger.error(f"Error extracting RTF text: {e}")
            return ""
    
    def _extract_odt_text(self, file_path: Path) -> str:
        """Извлечение текста из ODT"""
        if not odf_load:
            logger.warning("odfpy not available")
            return ""
        
        try:
            doc = odf_load(file_path)
            text = ""
            
            for paragraph in doc.getElementsByType(text.P):
                text += teletype.extractText(paragraph) + "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting ODT text: {e}")
            return ""
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """Извлечение текста из TXT"""
        try:
            # Пробуем разные кодировки
            encodings = ['utf-8', 'cp1251', 'latin1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            logger.warning(f"Could not decode text file {file_path}")
            return ""
            
        except Exception as e:
            logger.error(f"Error extracting TXT text: {e}")
            return ""
    
    def _extract_html_text(self, file_path: Path) -> str:
        """Извлечение текста из HTML"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Простое удаление HTML тегов
            text = re.sub(r'<[^>]+>', '', content)
            text = re.sub(r'&[a-zA-Z]+;', ' ', text)  # HTML entities
            
            return text
        except Exception as e:
            logger.error(f"Error extracting HTML text: {e}")
            return ""
    
    def _analyze_content(self, result: TemplateAnalysisResult, text_content: str):
        """Анализ содержимого документа"""
        lines = text_content.split('\n')
        
        # Пытаемся найти заголовок документа
        for line in lines[:10]:  # Смотрим первые 10 строк
            line = line.strip()
            if len(line) > 10 and len(line) < 100:
                # Проверяем, что это не плейсхолдер
                if not re.search(r'[{}\[\]_]{3,}', line):
                    result.title = line
                    break
        
        # Создаем описание из первых нескольких строк
        description_lines = []
        for line in lines[:20]:
            line = line.strip()
            if line and len(line) > 20:
                description_lines.append(line)
                if len(description_lines) >= 3:
                    break
        
        result.description = ' '.join(description_lines)[:500]  # Ограничиваем длину
    
    def _categorize_document(self, text_content: str) -> DocumentCategory:
        """Определение категории документа на основе содержимого"""
        text_lower = text_content.lower()
        
        category_scores = {}
        
        for category, patterns in self.category_patterns.items():
            score = 0
            for pattern in patterns:
                matches = len(re.findall(pattern, text_lower))
                score += matches
            
            if score > 0:
                category_scores[category] = score
        
        if category_scores:
            # Возвращаем категорию с наибольшим счетом
            return max(category_scores, key=category_scores.get)
        
        return DocumentCategory.TEMPLATE  # По умолчанию    

    def _analyze_structure(self, text_content: str) -> TemplateStructure:
        """Анализ структуры документа"""
        structure = TemplateStructure()
        
        lines = text_content.split('\n')
        structure.pages = max(1, text_content.count('\f'))  # Form feed для страниц
        structure.paragraphs = len([line for line in lines if line.strip()])
        
        # Поиск заголовков
        for line in lines:
            line = line.strip()
            if line:
                for pattern in self.structure_patterns['headers']:
                    if re.match(pattern, line):
                        structure.headers.append(line)
                        break
        
        # Поиск разделов
        for line in lines:
            line = line.strip().lower()
            for pattern in self.structure_patterns['sections']:
                if re.search(pattern, line):
                    structure.sections.append(line)
        
        # Подсчет таблиц (упрощенная эвристика)
        table_indicators = ['\t', '|', '┌', '┐', '└', '┘', '├', '┤']
        table_lines = 0
        for line in lines:
            if any(indicator in line for indicator in table_indicators):
                table_lines += 1
        
        structure.tables = table_lines // 3  # Примерная оценка
        
        # Подсчет списков
        list_patterns = [r'^\s*[-*•]\s', r'^\s*\d+\.\s', r'^\s*[a-zA-Z]\)\s']
        for line in lines:
            for pattern in list_patterns:
                if re.match(pattern, line):
                    structure.lists += 1
                    break
        
        return structure
    
    def _extract_placeholders(self, text_content: str) -> List[PlaceholderInfo]:
        """Извлечение плейсхолдеров из текста"""
        placeholders = []
        found_placeholders = set()  # Для избежания дубликатов
        
        for pattern, pattern_type in self.placeholder_patterns:
            matches = re.finditer(pattern, text_content, re.IGNORECASE)
            
            for match in matches:
                placeholder_text = match.group(0)
                
                if placeholder_text in found_placeholders:
                    continue
                
                found_placeholders.add(placeholder_text)
                
                # Определяем тип плейсхолдера
                placeholder_type = self._determine_placeholder_type(placeholder_text)
                
                # Извлекаем имя плейсхолдера
                if len(match.groups()) > 0:
                    name = match.group(1).strip()
                else:
                    name = placeholder_text
                
                placeholder = PlaceholderInfo(
                    name=name,
                    type=placeholder_type,
                    pattern=placeholder_text,
                    description=self._generate_placeholder_description(name, placeholder_type)
                )
                
                placeholders.append(placeholder)
        
        return placeholders
    
    def _determine_placeholder_type(self, placeholder_text: str) -> str:
        """Определение типа плейсхолдера"""
        text_lower = placeholder_text.lower()
        
        # Типы на основе ключевых слов
        type_keywords = {
            'date': ['дата', 'число', 'date', 'день', 'месяц', 'год'],
            'name': ['фио', 'имя', 'фамилия', 'name', 'ф.и.о'],
            'email': ['email', 'почта', 'e-mail', 'электронная'],
            'phone': ['телефон', 'тел', 'phone', 'моб'],
            'address': ['адрес', 'address', 'место'],
            'number': ['номер', 'num', 'number', '№'],
            'money': ['сумма', 'стоимость', 'цена', 'руб', 'amount'],
            'signature': ['подпись', 'signature', 'sign']
        }
        
        for field_type, keywords in type_keywords.items():
            if any(keyword in text_lower for keyword in keywords):
                return field_type
        
        # Проверяем паттерны
        if re.search(r'_{3,}', placeholder_text):
            return 'text'
        elif re.search(r'\.{3,}', placeholder_text):
            return 'text'
        
        return 'text'  # По умолчанию
    
    def _generate_placeholder_description(self, name: str, placeholder_type: str) -> str:
        """Генерация описания для плейсхолдера"""
        descriptions = {
            'date': 'Дата в формате ДД.ММ.ГГГГ',
            'name': 'Полное имя (Фамилия Имя Отчество)',
            'email': 'Адрес электронной почты',
            'phone': 'Номер телефона',
            'address': 'Почтовый адрес',
            'number': 'Числовое значение',
            'money': 'Денежная сумма',
            'signature': 'Место для подписи',
            'text': 'Текстовое поле'
        }
        
        base_description = descriptions.get(placeholder_type, 'Поле для ввода')
        
        if name and name != placeholder_type:
            return f"{name} - {base_description}"
        
        return base_description
    
    def _detect_language(self, text_content: str) -> str:
        """Определение языка документа"""
        text_lower = text_content.lower()
        words = re.findall(r'\b[а-яёa-z]+\b', text_lower)
        
        if not words:
            return "unknown"
        
        language_scores = {}
        
        for lang, keywords in self.language_keywords.items():
            score = sum(1 for word in words if word in keywords)
            language_scores[lang] = score
        
        if language_scores:
            detected_lang = max(language_scores, key=language_scores.get)
            # Проверяем, что счет достаточно высокий
            if language_scores[detected_lang] > len(words) * 0.01:  # Минимум 1% ключевых слов
                return detected_lang
        
        # Дополнительная проверка по алфавиту
        cyrillic_chars = len(re.findall(r'[а-яё]', text_lower))
        latin_chars = len(re.findall(r'[a-z]', text_lower))
        
        if cyrillic_chars > latin_chars:
            return "ru"
        elif latin_chars > cyrillic_chars:
            return "en"
        
        return "unknown"    

    def _extract_keywords(self, text_content: str) -> List[str]:
        """Извлечение ключевых слов из текста"""
        # Простая реализация - в реальности можно использовать NLP библиотеки
        text_lower = text_content.lower()
        
        # Удаляем стоп-слова и извлекаем значимые слова
        stop_words = {
            'ru': ['и', 'в', 'на', 'с', 'по', 'для', 'от', 'до', 'при', 'что', 'как', 'где', 'это', 'все', 'еще', 'уже', 'только', 'может', 'быть'],
            'en': ['the', 'and', 'of', 'to', 'in', 'for', 'with', 'on', 'at', 'by', 'from', 'this', 'that', 'all', 'any', 'can', 'will']
        }
        
        # Извлекаем слова длиной от 3 до 20 символов
        words = re.findall(r'\b[а-яёa-z]{3,20}\b', text_lower)
        
        # Подсчитываем частоту слов
        word_freq = {}
        for word in words:
            # Пропускаем стоп-слова
            if word not in stop_words.get('ru', []) and word not in stop_words.get('en', []):
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Возвращаем топ-20 самых частых слов
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return [word for word, freq in sorted_words[:20] if freq > 1]
    
    def _generate_tags(self, result: TemplateAnalysisResult) -> List[str]:
        """Генерация тегов для шаблона"""
        tags = []
        
        # Теги на основе категории
        tags.append(result.category.value)
        
        # Теги на основе формата
        tags.append(result.format.value)
        
        # Теги на основе языка
        if result.language != "unknown":
            tags.append(result.language)
        
        # Теги на основе сложности
        tags.append(result.complexity.value)
        
        # Теги на основе структуры
        if result.structure.tables > 0:
            tags.append("tables")
        
        if result.structure.lists > 0:
            tags.append("lists")
        
        if len(result.structure.headers) > 5:
            tags.append("structured")
        
        # Теги на основе плейсхолдеров
        if result.estimated_fields > 10:
            tags.append("many_fields")
        elif result.estimated_fields > 5:
            tags.append("medium_fields")
        elif result.estimated_fields > 0:
            tags.append("few_fields")
        
        # Теги на основе ключевых слов
        business_keywords = ['договор', 'контракт', 'соглашение', 'contract', 'agreement']
        if any(keyword in result.keywords for keyword in business_keywords):
            tags.append("business")
        
        legal_keywords = ['закон', 'право', 'юридический', 'legal', 'law']
        if any(keyword in result.keywords for keyword in legal_keywords):
            tags.append("legal")
        
        return list(set(tags))  # Удаляем дубликаты
    
    def _assess_complexity(self, result: TemplateAnalysisResult) -> ComplexityLevel:
        """Оценка сложности шаблона"""
        complexity_score = 0
        
        # Факторы сложности
        complexity_score += len(result.placeholders) * 2
        complexity_score += result.structure.tables * 5
        complexity_score += result.structure.lists * 2
        complexity_score += len(result.structure.headers) * 1
        complexity_score += len(result.structure.sections) * 3
        complexity_score += result.structure.pages * 2
        
        # Дополнительные факторы
        if result.file_size > 1024 * 1024:  # Больше 1 MB
            complexity_score += 10
        
        if len(result.keywords) > 50:
            complexity_score += 5
        
        # Определяем уровень сложности
        if complexity_score < 10:
            return ComplexityLevel.SIMPLE
        elif complexity_score < 30:
            return ComplexityLevel.MEDIUM
        elif complexity_score < 60:
            return ComplexityLevel.COMPLEX
        else:
            return ComplexityLevel.VERY_COMPLEX
    
    def _calculate_quality_scores(self, result: TemplateAnalysisResult, text_content: str):
        """Вычисление оценок качества шаблона"""
        
        # Оценка качества (0.0 - 1.0)
        quality_factors = []
        
        # Наличие заголовка
        if result.title:
            quality_factors.append(0.2)
        
        # Наличие структуры
        if len(result.structure.headers) > 0:
            quality_factors.append(0.15)
        
        # Наличие плейсхолдеров
        if len(result.placeholders) > 0:
            quality_factors.append(0.25)
        
        # Читаемость текста
        if len(text_content.strip()) > 100:
            quality_factors.append(0.1)
        
        # Отсутствие ошибок
        if len(result.errors) == 0:
            quality_factors.append(0.1)
        
        # Наличие описания
        if result.description:
            quality_factors.append(0.1)
        
        # Определенный язык
        if result.language != "unknown":
            quality_factors.append(0.1)
        
        result.quality_score = sum(quality_factors)
        
        # Оценка полноты (0.0 - 1.0)
        completeness_factors = []
        
        # Все основные поля заполнены
        if result.title and result.description:
            completeness_factors.append(0.3)
        
        # Есть плейсхолдеры
        if len(result.placeholders) > 0:
            completeness_factors.append(0.3)
        
        # Есть структура
        if len(result.structure.headers) > 0 or len(result.structure.sections) > 0:
            completeness_factors.append(0.2)
        
        # Есть ключевые слова
        if len(result.keywords) > 0:
            completeness_factors.append(0.1)
        
        # Есть теги
        if len(result.tags) > 0:
            completeness_factors.append(0.1)
        
        result.completeness_score = sum(completeness_factors)
        
        # Оценка удобства использования (0.0 - 1.0)
        usability_factors = []
        
        # Понятные плейсхолдеры
        clear_placeholders = sum(1 for p in result.placeholders if p.description)
        if clear_placeholders > 0:
            usability_factors.append(min(0.4, clear_placeholders / len(result.placeholders) * 0.4))
        
        # Умеренная сложность
        if result.complexity in [ComplexityLevel.SIMPLE, ComplexityLevel.MEDIUM]:
            usability_factors.append(0.3)
        
        # Хорошая структура
        if len(result.structure.headers) > 0:
            usability_factors.append(0.2)
        
        # Поддерживаемый формат
        if result.format in [DocumentFormat.PDF, DocumentFormat.DOCX, DocumentFormat.TXT]:
            usability_factors.append(0.1)
        
        result.usability_score = sum(usability_factors)
    
    def _extract_file_metadata(self, result: TemplateAnalysisResult, file_path: Path, document_format: DocumentFormat):
        """Извлечение метаданных файла"""
        try:
            stat = file_path.stat()
            result.creation_date = datetime.fromtimestamp(stat.st_ctime)
            result.modification_date = datetime.fromtimestamp(stat.st_mtime)
            
            # Для DOCX можем извлечь дополнительные метаданные
            if document_format == DocumentFormat.DOCX and docx:
                try:
                    doc = docx.Document(file_path)
                    core_props = doc.core_properties
                    
                    if core_props.author:
                        result.author = core_props.author
                    
                    if core_props.title:
                        result.title = result.title or core_props.title
                    
                    if core_props.created:
                        result.creation_date = core_props.created
                    
                    if core_props.modified:
                        result.modification_date = core_props.modified
                        
                except Exception as e:
                    logger.warning(f"Could not extract DOCX metadata: {e}")
                    
        except Exception as e:
            logger.error(f"Error extracting file metadata: {e}")
            result.errors.append(f"Metadata extraction error: {e}")


# Функции для интеграции с унифицированной системой
def analyze_template_file(file_path: str, **kwargs) -> Dict[str, Any]:
    """
    Анализ файла шаблона
    
    Args:
        file_path: Путь к файлу шаблона
        **kwargs: Дополнительные параметры
    """
    try:
        analyzer = AdvancedTemplateAnalyzer()
        result = analyzer.analyze_template(file_path)
        
        # Преобразуем результат в словарь для JSON сериализации
        return {
            'status': 'success',
            'file_path': result.file_path,
            'analysis': {
                'format': result.format.value,
                'category': result.category.value,
                'complexity': result.complexity.value,
                'title': result.title,
                'description': result.description,
                'language': result.language,
                'file_size': result.file_size,
                'structure': {
                    'sections': len(result.structure.sections),
                    'headers': len(result.structure.headers),
                    'tables': result.structure.tables,
                    'lists': result.structure.lists,
                    'pages': result.structure.pages,
                    'paragraphs': result.structure.paragraphs
                },
                'placeholders': [
                    {
                        'name': p.name,
                        'type': p.type,
                        'pattern': p.pattern,
                        'required': p.required,
                        'description': p.description
                    }
                    for p in result.placeholders
                ],
                'estimated_fields': result.estimated_fields,
                'quality_scores': {
                    'quality': result.quality_score,
                    'completeness': result.completeness_score,
                    'usability': result.usability_score
                },
                'tags': result.tags,
                'keywords': result.keywords[:10],  # Топ-10 ключевых слов
                'author': result.author,
                'organization': result.organization,
                'creation_date': result.creation_date.isoformat() if result.creation_date else None,
                'modification_date': result.modification_date.isoformat() if result.modification_date else None,
                'processing_time': result.processing_time,
                'errors': result.errors,
                'warnings': result.warnings
            }
        }
        
    except Exception as e:
        logger.error(f"Error analyzing template: {e}")
        return {
            'status': 'error',
            'error': str(e),
            'file_path': file_path
        }

# Регистрация в унифицированной системе инструментов
TEMPLATE_ANALYSIS_TOOLS = {
    'analyze_template_file': {
        'function': analyze_template_file,
        'description': 'Анализ файла шаблона с извлечением метаданных, структуры и плейсхолдеров',
        'category': 'document_analysis',
        'ui_placement': 'tools',
        'parameters': {
            'file_path': 'str - Путь к файлу шаблона для анализа'
        }
    }
}

if __name__ == "__main__":
    # Тестирование анализатора шаблонов
    analyzer = AdvancedTemplateAnalyzer()
    
    # Создаем тестовый файл
    test_content = """
    ДОГОВОР СТРОИТЕЛЬНОГО ПОДРЯДА
    
    г. {ГОРОД}, {ДАТА}
    
    {ЗАКАЗЧИК}, именуемый в дальнейшем "Заказчик", с одной стороны, и {ПОДРЯДЧИК}, 
    именуемый в дальнейшем "Подрядчик", с другой стороны, заключили настоящий договор о нижеследующем:
    
    1. ПРЕДМЕТ ДОГОВОРА
    1.1. Подрядчик обязуется выполнить строительные работы по адресу: {АДРЕС_ОБЪЕКТА}
    1.2. Стоимость работ составляет: {СУММА} рублей
    
    2. СРОКИ ВЫПОЛНЕНИЯ
    2.1. Начало работ: {ДАТА_НАЧАЛА}
    2.2. Окончание работ: {ДАТА_ОКОНЧАНИЯ}
    
    Подписи сторон:
    
    Заказчик: _________________ {ФИО_ЗАКАЗЧИКА}
    
    Подрядчик: ________________ {ФИО_ПОДРЯДЧИКА}
    """
    
    # Сохраняем тестовый файл
    test_file = Path("test_template.txt")
    test_file.write_text(test_content, encoding='utf-8')
    
    try:
        # Анализируем тестовый файл
        result = analyzer.analyze_template(str(test_file))
        
        print("=== РЕЗУЛЬТАТ АНАЛИЗА ШАБЛОНА ===")
        print(f"Файл: {result.file_path}")
        print(f"Формат: {result.format.value}")
        print(f"Категория: {result.category.value}")
        print(f"Сложность: {result.complexity.value}")
        print(f"Заголовок: {result.title}")
        print(f"Язык: {result.language}")
        print(f"Найдено плейсхолдеров: {len(result.placeholders)}")
        
        print("\nПлейсхолдеры:")
        for placeholder in result.placeholders[:5]:  # Показываем первые 5
            print(f"  - {placeholder.name} ({placeholder.type}): {placeholder.description}")
        
        print(f"\nОценки качества:")
        print(f"  Качество: {result.quality_score:.2f}")
        print(f"  Полнота: {result.completeness_score:.2f}")
        print(f"  Удобство: {result.usability_score:.2f}")
        
        print(f"\nТеги: {', '.join(result.tags)}")
        print(f"Ключевые слова: {', '.join(result.keywords[:5])}")
        
        if result.errors:
            print(f"\nОшибки: {result.errors}")
        
        print(f"\nВремя обработки: {result.processing_time:.3f} сек")
        
    finally:
        # Удаляем тестовый файл
        if test_file.exists():
            test_file.unlink()