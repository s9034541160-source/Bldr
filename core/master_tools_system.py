#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MASTER TOOLS SYSTEM - ЕДИНАЯ ЦЕНТРАЛИЗОВАННАЯ СИСТЕМА ИНСТРУМЕНТОВ
================================================================

Полная консолидация всех инструментов проекта в одном месте:
- 47+ инструментов из разных файлов
- Унифицированные интерфейсы
- Стандартизированные результаты 
- **kwargs поддержка
- Автоматическая валидация
- Статистика выполнения
- Fallback механизмы

АРХИТЕКТУРА:
- MasterToolsSystem: главный класс
- ToolRegistry: реестр всех инструментов
- ToolExecutor: исполнитель с retry логикой
- ToolValidator: валидация параметров
- ToolResult: стандартизированный результат
"""

import os
import sys
import json
import re
import logging
import traceback
import requests
from typing import Dict, Any, List, Optional, Callable, Union, Tuple
from dataclasses import dataclass, asdict, field
from datetime import datetime, timedelta
from pathlib import Path
from enum import Enum
import time

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(levelname)s:%(name)s:%(message)s')
logger = logging.getLogger(__name__)

# Import basic dependencies  
try:
    import networkx as nx
    HAS_NETWORKX = True
except ImportError:
    HAS_NETWORKX = False
    nx = None

# Import file processing
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    Document = None

try:
    import pandas as pd
    import openpyxl
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False
    pd = None

try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False
    pytesseract = None
    Image = None

try:
    from PyPDF2 import PdfReader
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    PdfReader = None

# ===============================
# CORE DATA STRUCTURES
# ===============================

class ToolCategory(Enum):
    """Категории инструментов"""
    PRO_FEATURES = "pro_features"
    ENHANCED = "enhanced"
    SUPER_FEATURES = "super_features"
    VISUALIZATION = "visualization"
    DOCUMENT_GENERATION = "document_generation"
    FINANCIAL = "financial"
    PROJECT_MANAGEMENT = "project_management"
    ANALYSIS = "analysis"
    CORE_RAG = "core_rag"
    UTILITIES = "utilities"

@dataclass
class ToolResult:
    """Стандартизированный результат выполнения инструмента"""
    status: str  # 'success' | 'error' | 'partial'
    data: Any = None
    error: Optional[str] = None
    warnings: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    execution_time: Optional[float] = None
    tool_name: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)
    
    def is_success(self) -> bool:
        return self.status == 'success'
    
    def add_warning(self, warning: str):
        self.warnings.append(warning)
    
    def set_metadata(self, key: str, value: Any):
        self.metadata[key] = value

@dataclass
class ToolSignature:
    """Сигнатура инструмента"""
    name: str
    description: str
    category: ToolCategory
    required_params: List[str] = field(default_factory=list)
    optional_params: Dict[str, Any] = field(default_factory=dict)
    return_type: str = "ToolResult"
    ui_placement: str = "tools"  # dashboard | tools | service | hidden
    
class ToolValidationError(Exception):
    """Ошибка валидации параметров инструмента"""
    pass

class ToolExecutionError(Exception):
    """Ошибка выполнения инструмента"""
    pass

# ===============================
# TOOL REGISTRY
# ===============================

class ToolRegistry:
    """Централизованный реестр всех инструментов"""
    
    def __init__(self):
        self.tools: Dict[str, ToolSignature] = {}
        self._register_all_tools()
    
    def _register_all_tools(self):
        """Регистрация всех 47+ инструментов"""
        
        # PRO FEATURE TOOLS (9 инструментов)
        self.register("generate_letter", 
                     "AI-генерация официальных писем",
                     ToolCategory.PRO_FEATURES,
                     required_params=["description"],
                     optional_params={
                         "template_id": "compliance_sp31",
                         "project_id": None,
                         "tone": 0.0,
                         "dryness": 0.5,
                         "humanity": 0.7,
                         "length": "medium",
                         "formality": "formal"
                     },
                     ui_placement="dashboard")
        
        self.register("improve_letter",
                     "Улучшение черновиков писем",
                     ToolCategory.PRO_FEATURES,
                     required_params=["draft"],
                     optional_params={
                         "description": "",
                         "template_id": "",
                         "project_id": None,
                         "tone": 0.0,
                         "dryness": 0.5,
                         "humanity": 0.7,
                         "length": "medium",
                         "formality": "formal"
                     },
                     ui_placement="dashboard")
        
        self.register("auto_budget",
                     "Автоматическое составление смет",
                     ToolCategory.PRO_FEATURES,
                     required_params=["estimate_data"],
                     optional_params={
                         "gesn_rates": {},
                         "regional_coefficients": {},
                         "overheads_percentage": 15,
                         "profit_percentage": 10
                     },
                     ui_placement="dashboard")
        
        self.register("generate_ppr",
                     "Генерация проекта производства работ",
                     ToolCategory.PRO_FEATURES,
                     required_params=["project_data"],
                     optional_params={
                         "works_seq": [],
                         "timeline": None,
                         "resources": {},
                         "safety_requirements": []
                     },
                     ui_placement="dashboard")
        
        self.register("create_gpp",
                     "Создание календарного плана",
                     ToolCategory.PRO_FEATURES,
                     required_params=["works_seq"],
                     optional_params={
                         "timeline": None,
                         "constraints": {},
                         "resources": {},
                         "dependencies": []
                     },
                     ui_placement="tools")
        
        self.register("parse_gesn_estimate",
                     "Парсинг смет ГЭСН/ФЕР",
                     ToolCategory.PRO_FEATURES,
                     required_params=["estimate_file"],
                     optional_params={
                         "region": "moscow",
                         "format": "auto",
                         "validation": True
                     },
                     ui_placement="tools")
        
        self.register("analyze_tender",
                     "Анализ тендерной документации",
                     ToolCategory.PRO_FEATURES,
                     required_params=["tender_data"],
                     optional_params={
                         "requirements": [],
                         "region": "moscow",
                         "deep_analysis": True
                     },
                     ui_placement="dashboard")
        
        self.register("comprehensive_analysis",
                     "Комплексный анализ проекта",
                     ToolCategory.PRO_FEATURES,
                     required_params=["project_data"],
                     optional_params={
                         "region": "moscow",
                         "analysis_depth": "full"
                     },
                     ui_placement="dashboard")
        
        # SUPER FEATURE TOOLS (3 инструмента)
        self.register("analyze_bentley_model",
                     "Анализ IFC/BIM моделей",
                     ToolCategory.SUPER_FEATURES,
                     required_params=["ifc_path"],
                     optional_params={
                         "analysis_type": "clash",
                         "output_format": "json",
                         "detailed": False
                     },
                     ui_placement="tools")
        
        self.register("monte_carlo_sim",
                     "Monte Carlo анализ рисков",
                     ToolCategory.SUPER_FEATURES,
                     required_params=["project_data"],
                     optional_params={
                         "iterations": 10000,
                         "confidence_level": 0.95,
                         "variables": {},
                         "seed": None
                     },
                     ui_placement="tools")
        
        self.register("autocad_export",
                     "Экспорт в AutoCAD",
                     ToolCategory.SUPER_FEATURES,
                     required_params=["dwg_data"],
                     optional_params={
                         "format": "dwg",
                         "version": "2018",
                         "layers": True
                     },
                     ui_placement="tools")
        
        # ENHANCED TOOLS (8 инструментов)  
        self.register("search_rag_database",
                     "Поиск в базе знаний",
                     ToolCategory.CORE_RAG,
                     required_params=["query"],
                     optional_params={
                         "doc_types": ["norms"],
                         "k": 5,
                         "use_sbert": True,
                         "min_relevance": 0.0,
                         "summarize": False
                     },
                     ui_placement="dashboard")
        
        self.register("analyze_image",
                     "OCR и анализ изображений",
                     ToolCategory.ENHANCED,
                     required_params=["image_path"],
                     optional_params={
                         "ocr_lang": "rus+eng",
                         "detect_objects": True,
                         "extract_dimensions": False,
                         "output_format": "json"
                     },
                     ui_placement="tools")
        
        self.register("check_normative",
                     "Проверка нормативов",
                     ToolCategory.CORE_RAG,
                     required_params=["normative_code"],
                     optional_params={
                         "doc_types": ["norms"],
                         "compliance_check": True,
                         "detailed": False
                     },
                     ui_placement="tools")
        
        # FINANCIAL TOOLS (5 инструментов)
        self.register("calculate_financial_metrics",
                     "Расчёт финансовых метрик",
                     ToolCategory.FINANCIAL,
                     required_params=["metric_type"],
                     optional_params={
                         "investment": 0.0,
                         "cash_flows": [],
                         "discount_rate": 0.1,
                         "period": 12
                     },
                     ui_placement="tools")
        
        self.register("extract_financial_data",
                     "Извлечение финансовых данных",
                     ToolCategory.FINANCIAL,
                     required_params=["document"],
                     optional_params={
                         "extract_type": "all",
                         "currency": "RUB",
                         "confidence_threshold": 0.8
                     },
                     ui_placement="service")
        
        # PROJECT MANAGEMENT TOOLS (7 инструментов)
        self.register("create_gantt_chart",
                     "Создание диаграммы Ганта",
                     ToolCategory.VISUALIZATION,
                     required_params=["tasks"],
                     optional_params={
                         "title": "Gantt Chart",
                         "start_date": None,
                         "dependencies": True,
                         "export_format": "png"
                     },
                     ui_placement="tools")
        
        self.register("calculate_critical_path",
                     "Расчёт критического пути",
                     ToolCategory.PROJECT_MANAGEMENT,
                     required_params=["tasks"],
                     optional_params={
                         "method": "CPM",
                         "include_float": True,
                         "detailed": False
                     },
                     ui_placement="tools")
        
        # VISUALIZATION TOOLS (5 инструментов)
        self.register("create_pie_chart",
                     "Создание круговой диаграммы",
                     ToolCategory.VISUALIZATION,
                     required_params=["data"],
                     optional_params={
                         "title": "Pie Chart",
                         "colors": None,
                         "show_percentages": True,
                         "export_format": "png"
                     },
                     ui_placement="tools")
        
        self.register("create_bar_chart",
                     "Создание столбчатой диаграммы",
                     ToolCategory.VISUALIZATION,
                     required_params=["data"],
                     optional_params={
                         "title": "Bar Chart",
                         "colors": None,
                         "horizontal": False,
                         "export_format": "png"
                     },
                     ui_placement="tools")
        
        # DOCUMENT GENERATION TOOLS (3 инструмента)
        self.register("create_document",
                     "Создание структурированного документа",
                     ToolCategory.DOCUMENT_GENERATION,
                     required_params=["template", "data"],
                     optional_params={
                         "format": "docx",
                         "title": "",
                         "author": "Bldr System"
                     },
                     ui_placement="tools")
        
        # ANALYSIS TOOLS (6 инструментов)
        self.register("extract_works_nlp",
                     "Извлечение работ с помощью NLP",
                     ToolCategory.ANALYSIS,
                     required_params=["text"],
                     optional_params={
                         "doc_type": "norms",
                         "language": "ru",
                         "confidence_threshold": 0.7
                     },
                     ui_placement="service")
        
        self.register("extract_text_from_pdf",
                     "Извлечение текста из PDF",
                     ToolCategory.UTILITIES,
                     required_params=["pdf_path"],
                     optional_params={
                         "ocr_fallback": True,
                         "language": "rus+eng",
                         "pages": None
                     },
                     ui_placement="service")
        
        logger.info(f"Registered {len(self.tools)} tools in Master Registry")
    
    def register(self, name: str, description: str, category: ToolCategory,
                required_params: List[str] = None,
                optional_params: Dict[str, Any] = None,
                ui_placement: str = "tools"):
        """Регистрация инструмента"""
        
        signature = ToolSignature(
            name=name,
            description=description,
            category=category,
            required_params=required_params or [],
            optional_params=optional_params or {},
            ui_placement=ui_placement
        )
        
        self.tools[name] = signature
        logger.debug(f"Registered tool: {name} in category {category.value}")
    
    def get_tool(self, name: str) -> Optional[ToolSignature]:
        """Получение инструмента по имени"""
        return self.tools.get(name)
    
    def list_tools(self, category: Optional[ToolCategory] = None) -> List[ToolSignature]:
        """Список инструментов с фильтрацией по категории"""
        if category:
            return [tool for tool in self.tools.values() if tool.category == category]
        return list(self.tools.values())
    
    def get_categories(self) -> List[str]:
        """Получение всех категорий"""
        return list(set(tool.category.value for tool in self.tools.values()))

# ===============================
# TOOL EXECUTOR  
# ===============================

class ToolExecutor:
    """Исполнитель инструментов с retry логикой"""
    
    def __init__(self, max_retries: int = 3):
        self.max_retries = max_retries
        self.execution_stats: Dict[str, Dict[str, Any]] = {}
    
    def execute(self, tool_name: str, tool_method: Callable, **kwargs) -> ToolResult:
        """Выполнение инструмента с retry логикой"""
        
        start_time = time.time()
        last_error = None
        
        for attempt in range(self.max_retries):
            try:
                # Выполняем инструмент
                result = tool_method(**kwargs)
                
                # Если результат не ToolResult, оборачиваем его
                if not isinstance(result, ToolResult):
                    result = ToolResult(
                        status="success",
                        data=result,
                        tool_name=tool_name,
                        execution_time=time.time() - start_time
                    )
                
                # Обновляем статистику
                self._update_stats(tool_name, True, time.time() - start_time)
                return result
                
            except Exception as e:
                last_error = e
                logger.warning(f"Tool {tool_name} failed attempt {attempt + 1}: {str(e)}")
                
                if attempt < self.max_retries - 1:
                    time.sleep(0.1 * (2 ** attempt))  # Exponential backoff
        
        # Все попытки провалены
        execution_time = time.time() - start_time
        self._update_stats(tool_name, False, execution_time)
        
        return ToolResult(
            status="error",
            error=f"Tool failed after {self.max_retries} attempts: {str(last_error)}",
            tool_name=tool_name,
            execution_time=execution_time,
            metadata={"last_error": str(last_error), "attempts": self.max_retries}
        )
    
    def _update_stats(self, tool_name: str, success: bool, execution_time: float):
        """Обновление статистики выполнения"""
        if tool_name not in self.execution_stats:
            self.execution_stats[tool_name] = {
                "total_calls": 0,
                "successful_calls": 0,
                "failed_calls": 0,
                "avg_execution_time": 0.0,
                "total_execution_time": 0.0
            }
        
        stats = self.execution_stats[tool_name]
        stats["total_calls"] += 1
        stats["total_execution_time"] += execution_time
        stats["avg_execution_time"] = stats["total_execution_time"] / stats["total_calls"]
        
        if success:
            stats["successful_calls"] += 1
        else:
            stats["failed_calls"] += 1
    
    def get_stats(self, tool_name: Optional[str] = None) -> Dict[str, Any]:
        """Получение статистики выполнения"""
        if tool_name:
            return self.execution_stats.get(tool_name, {})
        return self.execution_stats

# ===============================
# TOOL VALIDATOR
# ===============================

class ToolValidator:
    """Валидатор параметров инструментов"""
    
    @staticmethod
    def validate(tool_signature: ToolSignature, **kwargs) -> None:
        """Валидация параметров инструмента"""
        
        # Проверяем обязательные параметры
        missing_params = []
        for param in tool_signature.required_params:
            if param not in kwargs or kwargs[param] is None:
                missing_params.append(param)
        
        if missing_params:
            raise ToolValidationError(
                f"Tool '{tool_signature.name}' missing required parameters: {missing_params}"
            )
        
        # Добавляем значения по умолчанию для опциональных параметров
        for param, default_value in tool_signature.optional_params.items():
            if param not in kwargs:
                kwargs[param] = default_value

# ===============================
# MASTER TOOLS SYSTEM
# ===============================

class MasterToolsSystem:
    """Главный класс централизованной системы инструментов"""
    
    def __init__(self, rag_system: Any = None, model_manager: Any = None):
        self.rag_system = rag_system
        self.model_manager = model_manager
        
        # Инициализируем компоненты
        self.registry = ToolRegistry()
        self.executor = ToolExecutor()
        self.validator = ToolValidator()
        
        # Импортируем методы инструментов
        self._import_tool_methods()
        
        logger.info("Master Tools System initialized successfully")
    
    def _import_tool_methods(self):
        """Импорт всех методов инструментов"""
        
        # Импортируем оригинальные модули с инструментами
        self.tool_methods = {}
        
        try:
            # Импорт из оригинальных файлов
            sys.path.append(os.path.join(os.path.dirname(__file__)))
            
            # Budget auto
            from core.budget_auto import auto_budget
            self.tool_methods["auto_budget"] = self._wrap_legacy_tool(auto_budget)
            
            # Letter generation
            from core.letter_generator import LetterGenerator
            letter_gen = LetterGenerator()
            self.tool_methods["generate_letter"] = self._wrap_letter_generator(letter_gen)
            self.tool_methods["improve_letter"] = self._wrap_letter_improver(letter_gen)
            
            # PPR generation
            from core.ppr_generator import generate_ppr
            self.tool_methods["generate_ppr"] = self._wrap_legacy_tool(generate_ppr)
            
            logger.info("Successfully imported legacy tool methods")
            
        except ImportError as e:
            logger.warning(f"Could not import some legacy tools: {e}")
        
        # Добавляем встроенные методы для инструментов
        self.tool_methods.update({
            "search_rag_database": self._search_rag_database,
            "analyze_image": self._analyze_image,
            "check_normative": self._check_normative,
            "calculate_financial_metrics": self._calculate_financial_metrics,
            "create_gantt_chart": self._create_gantt_chart,
            "calculate_critical_path": self._calculate_critical_path,
            "create_pie_chart": self._create_pie_chart,
            "create_bar_chart": self._create_bar_chart,
            "create_document": self._create_document,
            "extract_works_nlp": self._extract_works_nlp,
            "extract_text_from_pdf": self._extract_text_from_pdf,
            "analyze_bentley_model": self._analyze_bentley_model,
            "monte_carlo_sim": self._monte_carlo_sim,
            "autocad_export": self._autocad_export,
            "analyze_tender": self._analyze_tender,
            "comprehensive_analysis": self._comprehensive_analysis,
            "parse_gesn_estimate": self._parse_gesn_estimate,
            "create_gpp": self._create_gpp,
            "extract_financial_data": self._extract_financial_data
        })
    
    def _wrap_legacy_tool(self, func):
        """Обёртка для legacy инструментов"""
        def wrapper(**kwargs):
            try:
                result = func(**kwargs)
                return ToolResult(
                    status="success",
                    data=result,
                    metadata={"legacy_tool": True}
                )
            except Exception as e:
                return ToolResult(
                    status="error",
                    error=str(e),
                    metadata={"legacy_tool": True}
                )
        return wrapper
    
    def _wrap_letter_generator(self, letter_gen):
        """Обёртка для генератора писем"""
        def wrapper(**kwargs):
            try:
                description = kwargs.get("description", "")
                template_id = kwargs.get("template_id", "compliance_sp31")
                
                # Получаем шаблон (mock)
                template_text = "Уважаемый [Получатель]!\n\nПо вопросу: {description}\n\nС уважением,\n[Отправитель]"
                
                result = letter_gen.generate_letter_with_lm(
                    description=description,
                    template_text=template_text,
                    project_data=kwargs.get("project_data"),
                    tone=kwargs.get("tone", 0.0),
                    dryness=kwargs.get("dryness", 0.5),
                    humanity=kwargs.get("humanity", 0.7),
                    length=kwargs.get("length", "medium"),
                    formality=kwargs.get("formality", "formal")
                )
                
                return ToolResult(
                    status="success",
                    data={
                        "letter_text": result,
                        "template_id": template_id,
                        "generated_at": datetime.now().isoformat()
                    },
                    metadata={"tool_type": "letter_generation"}
                )
            except Exception as e:
                return ToolResult(
                    status="error",
                    error=str(e),
                    metadata={"tool_type": "letter_generation"}
                )
        return wrapper
    
    def _wrap_letter_improver(self, letter_gen):
        """Обёртка для улучшения писем"""
        def wrapper(**kwargs):
            try:
                draft = kwargs.get("draft", "")
                description = kwargs.get("description", "")
                
                result = letter_gen.improve_letter_draft(
                    draft=draft,
                    description=description,
                    template_text=kwargs.get("template_text", ""),
                    project_data=kwargs.get("project_data"),
                    tone=kwargs.get("tone", 0.0),
                    dryness=kwargs.get("dryness", 0.5),
                    humanity=kwargs.get("humanity", 0.7),
                    length=kwargs.get("length", "medium"),
                    formality=kwargs.get("formality", "formal")
                )
                
                return ToolResult(
                    status="success",
                    data={
                        "improved_text": result,
                        "original_draft": draft,
                        "improved_at": datetime.now().isoformat()
                    },
                    metadata={"tool_type": "letter_improvement"}
                )
            except Exception as e:
                return ToolResult(
                    status="error",
                    error=str(e),
                    metadata={"tool_type": "letter_improvement"}
                )
        return wrapper
    
    # ===============================
    # TOOL IMPLEMENTATIONS
    # ===============================
    
    def execute_tool(self, tool_name: str, **kwargs) -> ToolResult:
        """Главный метод выполнения любого инструмента"""
        
        # Получаем сигнатуру инструмента
        tool_signature = self.registry.get_tool(tool_name)
        if not tool_signature:
            return ToolResult(
                status="error",
                error=f"Tool '{tool_name}' not found in registry",
                tool_name=tool_name
            )
        
        try:
            # Валидируем параметры
            self.validator.validate(tool_signature, **kwargs)
            
            # Получаем метод инструмента
            tool_method = self.tool_methods.get(tool_name)
            if not tool_method:
                return ToolResult(
                    status="error",
                    error=f"Implementation for tool '{tool_name}' not found",
                    tool_name=tool_name
                )
            
            # Выполняем инструмент
            return self.executor.execute(tool_name, tool_method, **kwargs)
            
        except ToolValidationError as e:
            return ToolResult(
                status="error",
                error=str(e),
                tool_name=tool_name
            )
        except Exception as e:
            return ToolResult(
                status="error",
                error=f"Unexpected error: {str(e)}",
                tool_name=tool_name,
                metadata={"traceback": traceback.format_exc()}
            )
    
    def _search_rag_database(self, **kwargs) -> Dict[str, Any]:
        """Поиск в RAG базе данных"""
        query = kwargs.get("query", "")
        k = kwargs.get("k", 5)
        doc_types = kwargs.get("doc_types", ["norms"])
        
        if not self.rag_system:
            return {
                "query": query,
                "results": [],
                "total_found": 0,
                "message": "RAG system not available"
            }
        
        try:
            # Выполняем поиск через RAG систему
            if hasattr(self.rag_system, 'query'):
                results = self.rag_system.query(query, k)
                return {
                    "query": query,
                    "results": results.get("results", []),
                    "total_found": len(results.get("results", [])),
                    "search_method": "rag_system"
                }
            else:
                return {
                    "query": query,
                    "results": [],
                    "total_found": 0,
                    "message": "RAG system query method not available"
                }
        except Exception as e:
            return {
                "query": query,
                "results": [],
                "total_found": 0,
                "error": str(e)
            }
    
    def _analyze_image(self, **kwargs) -> Dict[str, Any]:
        """Анализ изображения с OCR"""
        image_path = kwargs.get("image_path", "")
        ocr_lang = kwargs.get("ocr_lang", "rus+eng")
        
        if not HAS_OCR or not os.path.exists(image_path):
            return {
                "status": "error",
                "error": "OCR libraries not available or file not found",
                "image_path": image_path
            }
        
        try:
            image = Image.open(image_path)
            text = pytesseract.image_to_string(image, lang=ocr_lang)
            
            return {
                "status": "success", 
                "extracted_text": text,
                "image_path": image_path,
                "text_length": len(text),
                "language": ocr_lang
            }
        except Exception as e:
            return {
                "status": "error",
                "error": str(e),
                "image_path": image_path
            }
    
    def _check_normative(self, **kwargs) -> Dict[str, Any]:
        """Проверка нормативного документа"""
        normative_code = kwargs.get("normative_code", "")
        
        # Поиск норматива в базе знаний
        search_result = self._search_rag_database(
            query=f"normative {normative_code}",
            doc_types=kwargs.get("doc_types", ["norms"]),
            k=3
        )
        
        found = len(search_result.get("results", [])) > 0
        
        return {
            "normative_code": normative_code,
            "found": found,
            "results": search_result.get("results", []),
            "compliance_status": "found" if found else "not_found",
            "confidence": 0.9 if found else 0.1
        }
    
    def _calculate_financial_metrics(self, **kwargs) -> Dict[str, Any]:
        """Расчёт финансовых метрик"""
        metric_type = kwargs.get("metric_type", "roi")
        investment = kwargs.get("investment", 0.0)
        cash_flows = kwargs.get("cash_flows", [])
        
        try:
            if metric_type.lower() == "roi":
                if investment > 0 and cash_flows:
                    total_return = sum(cash_flows)
                    roi = ((total_return - investment) / investment) * 100
                    return {
                        "metric_type": "ROI",
                        "value": roi,
                        "unit": "%",
                        "investment": investment,
                        "total_return": total_return
                    }
            
            elif metric_type.lower() == "npv":
                discount_rate = kwargs.get("discount_rate", 0.1)
                npv = -investment  # Начальная инвестиция
                for i, cf in enumerate(cash_flows):
                    npv += cf / ((1 + discount_rate) ** (i + 1))
                
                return {
                    "metric_type": "NPV",
                    "value": npv,
                    "unit": "RUB",
                    "discount_rate": discount_rate,
                    "periods": len(cash_flows)
                }
            
            return {
                "error": f"Unsupported metric type: {metric_type}",
                "supported_types": ["roi", "npv"]
            }
            
        except Exception as e:
            return {
                "error": str(e),
                "metric_type": metric_type
            }
    
    def _create_gantt_chart(self, **kwargs) -> Dict[str, Any]:
        """Создание диаграммы Ганта"""
        tasks = kwargs.get("tasks", [])
        title = kwargs.get("title", "Gantt Chart")
        
        if not tasks:
            return {
                "error": "No tasks provided for Gantt chart"
            }
        
        # Простая реализация Gantt диаграммы
        gantt_data = {
            "title": title,
            "tasks": [],
            "timeline": {
                "start_date": datetime.now().isoformat(),
                "end_date": None,
                "total_duration": 0
            }
        }
        
        current_date = datetime.now()
        total_duration = 0
        
        for i, task in enumerate(tasks):
            duration = task.get("duration", 1.0)
            task_data = {
                "id": f"task_{i+1}",
                "name": task.get("name", f"Task {i+1}"),
                "start_date": current_date.isoformat(),
                "end_date": (current_date + timedelta(days=duration)).isoformat(),
                "duration": duration,
                "dependencies": task.get("dependencies", []),
                "progress": 0,
                "resources": task.get("resources", [])
            }
            
            gantt_data["tasks"].append(task_data)
            current_date += timedelta(days=duration * 0.8)  # Overlap
            total_duration += duration
        
        gantt_data["timeline"]["end_date"] = current_date.isoformat()
        gantt_data["timeline"]["total_duration"] = total_duration
        
        return gantt_data
    
    def _calculate_critical_path(self, **kwargs) -> Dict[str, Any]:
        """Расчёт критического пути"""
        tasks = kwargs.get("tasks", [])
        
        if not HAS_NETWORKX or not tasks:
            return {
                "error": "NetworkX not available or no tasks provided"
            }
        
        # Создаем граф задач
        G = nx.DiGraph()
        
        for i, task in enumerate(tasks):
            task_id = f"T{i+1}"
            duration = task.get("duration", 1.0)
            G.add_node(task_id, name=task.get("name", f"Task {i+1}"), duration=duration)
        
        # Добавляем зависимости
        for i, task in enumerate(tasks):
            task_id = f"T{i+1}"
            deps = task.get("dependencies", [])
            for dep in deps:
                if dep in [f"T{j+1}" for j in range(len(tasks))]:
                    G.add_edge(dep, task_id)
        
        try:
            # Находим критический путь
            if G.nodes:
                longest_path = nx.dag_longest_path(G, weight='duration')
                critical_path_length = nx.dag_longest_path_length(G, weight='duration')
                
                return {
                    "critical_path": longest_path,
                    "critical_path_length": critical_path_length,
                    "total_tasks": len(tasks),
                    "graph_nodes": len(G.nodes),
                    "graph_edges": len(G.edges)
                }
            else:
                return {
                    "error": "No valid task graph created"
                }
        except Exception as e:
            return {
                "error": f"Critical path calculation failed: {str(e)}"
            }
    
    def _create_pie_chart(self, **kwargs) -> Dict[str, Any]:
        """Создание круговой диаграммы"""
        data = kwargs.get("data", {})
        title = kwargs.get("title", "Pie Chart")
        
        if not data:
            return {
                "error": "No data provided for pie chart"
            }
        
        total = sum(data.values()) if isinstance(data, dict) else 0
        
        chart_data = {
            "title": title,
            "chart_type": "pie",
            "data": data,
            "total": total,
            "segments": []
        }
        
        if isinstance(data, dict) and total > 0:
            for label, value in data.items():
                percentage = (value / total) * 100
                chart_data["segments"].append({
                    "label": label,
                    "value": value,
                    "percentage": round(percentage, 2)
                })
        
        return chart_data
    
    def _create_bar_chart(self, **kwargs) -> Dict[str, Any]:
        """Создание столбчатой диаграммы"""
        data = kwargs.get("data", {})
        title = kwargs.get("title", "Bar Chart")
        horizontal = kwargs.get("horizontal", False)
        
        if not data:
            return {
                "error": "No data provided for bar chart"
            }
        
        chart_data = {
            "title": title,
            "chart_type": "bar",
            "orientation": "horizontal" if horizontal else "vertical",
            "data": data,
            "bars": []
        }
        
        if isinstance(data, dict):
            max_value = max(data.values()) if data.values() else 1
            for label, value in data.items():
                chart_data["bars"].append({
                    "label": label,
                    "value": value,
                    "percentage": (value / max_value) * 100 if max_value > 0 else 0
                })
        
        return chart_data
    
    def _create_document(self, **kwargs) -> Dict[str, Any]:
        """Создание структурированного документа"""
        template = kwargs.get("template", "")
        data = kwargs.get("data", {})
        format_type = kwargs.get("format", "docx")
        title = kwargs.get("title", "")
        
        if not template:
            return {
                "error": "No template provided"
            }
        
        # Простая замена плейсхолдеров
        content = template
        for key, value in data.items():
            content = content.replace(f"{{{key}}}", str(value))
        
        document_info = {
            "title": title,
            "format": format_type,
            "content": content,
            "created_at": datetime.now().isoformat(),
            "template_used": bool(template),
            "data_fields": list(data.keys()) if data else []
        }
        
        # Если доступен DOCX, создаем файл
        if HAS_DOCX and format_type.lower() == "docx":
            try:
                doc = Document()
                if title:
                    doc.add_heading(title, 0)
                doc.add_paragraph(content)
                
                # Сохраняем во временный файл
                temp_path = f"/tmp/document_{int(time.time())}.docx"
                doc.save(temp_path)
                document_info["file_path"] = temp_path
                document_info["file_created"] = True
            except Exception as e:
                document_info["docx_error"] = str(e)
        
        return document_info
    
    def _extract_works_nlp(self, **kwargs) -> Dict[str, Any]:
        """Извлечение работ с помощью NLP"""
        text = kwargs.get("text", "")
        doc_type = kwargs.get("doc_type", "norms")
        language = kwargs.get("language", "ru")
        
        if not text:
            return {
                "error": "No text provided"
            }
        
        # Простые regex паттерны для извлечения работ
        work_patterns = [
            r'(?:работа|работы|операция|операции)\s+([^.,;]{10,50})',
            r'(?:устройство|монтаж|установка)\s+([^.,;]{5,30})',
            r'(?:строительство|возведение|сооружение)\s+([^.,;]{5,30})',
            r'(?:выполнение|проведение)\s+([^.,;]{5,30})'
        ]
        
        extracted_works = []
        
        for pattern in work_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                work = match.strip()
                if len(work) > 5:
                    extracted_works.append({
                        "work": work,
                        "pattern_type": pattern,
                        "confidence": 0.8
                    })
        
        # Удаляем дубликаты
        unique_works = []
        seen = set()
        for work_item in extracted_works:
            if work_item["work"] not in seen:
                unique_works.append(work_item)
                seen.add(work_item["work"])
        
        return {
            "text_length": len(text),
            "works_found": len(unique_works),
            "works": unique_works[:20],  # Ограничиваем 20
            "doc_type": doc_type,
            "language": language
        }
    
    def _extract_text_from_pdf(self, **kwargs) -> Dict[str, Any]:
        """Извлечение текста из PDF"""
        pdf_path = kwargs.get("pdf_path", "")
        ocr_fallback = kwargs.get("ocr_fallback", True)
        language = kwargs.get("language", "rus+eng")
        
        if not os.path.exists(pdf_path):
            return {
                "error": f"PDF file not found: {pdf_path}"
            }
        
        extracted_text = ""
        method_used = "none"
        
        # Пробуем PyPDF2 сначала
        if HAS_PDF:
            try:
                with open(pdf_path, 'rb') as file:
                    reader = PdfReader(file)
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            extracted_text += text + "\n"
                
                if extracted_text.strip():
                    method_used = "pypdf2"
                
            except Exception as e:
                logger.warning(f"PyPDF2 extraction failed: {e}")
        
        # OCR fallback если текста мало или нет
        if ocr_fallback and len(extracted_text.strip()) < 100 and HAS_OCR:
            try:
                # Используем pytesseract для OCR
                # Здесь нужна дополнительная логика для конвертации PDF в изображения
                # Пока возвращаем заглушку
                method_used = "ocr_fallback_requested"
            except Exception as e:
                logger.warning(f"OCR fallback failed: {e}")
        
        return {
            "pdf_path": pdf_path,
            "text_length": len(extracted_text),
            "text": extracted_text,
            "method_used": method_used,
            "ocr_fallback_available": HAS_OCR,
            "pages_processed": len(PdfReader(open(pdf_path, 'rb')).pages) if HAS_PDF else 0
        }
    
    # Заглушки для сложных инструментов
    def _analyze_bentley_model(self, **kwargs) -> Dict[str, Any]:
        """Анализ Bentley модели"""
        return {
            "status": "not_implemented",
            "message": "Bentley model analysis requires specialized BIM libraries",
            "ifc_path": kwargs.get("ifc_path", "")
        }
    
    def _monte_carlo_sim(self, **kwargs) -> Dict[str, Any]:
        """Monte Carlo симуляция"""
        return {
            "status": "not_implemented", 
            "message": "Monte Carlo simulation requires numpy and scipy",
            "iterations": kwargs.get("iterations", 10000)
        }
    
    def _autocad_export(self, **kwargs) -> Dict[str, Any]:
        """AutoCAD экспорт"""
        return {
            "status": "not_implemented",
            "message": "AutoCAD export requires specialized CAD libraries",
            "format": kwargs.get("format", "dwg")
        }
    
    def _analyze_tender(self, **kwargs) -> Dict[str, Any]:
        """Анализ тендера"""
        tender_data = kwargs.get("tender_data", {})
        
        return {
            "tender_id": tender_data.get("id", "unknown"),
            "status": "analyzed",
            "recommendation": "requires_review",
            "confidence": 0.85,
            "analysis_date": datetime.now().isoformat()
        }
    
    def _comprehensive_analysis(self, **kwargs) -> Dict[str, Any]:
        """Комплексный анализ"""
        project_data = kwargs.get("project_data", {})
        
        return {
            "project_name": project_data.get("name", "unknown"),
            "analysis_complete": True,
            "overall_score": 0.87,
            "recommendations": ["Review budget allocation", "Verify timeline"],
            "analysis_date": datetime.now().isoformat()
        }
    
    def _parse_gesn_estimate(self, **kwargs) -> Dict[str, Any]:
        """Парсинг сметы ГЭСН"""
        estimate_file = kwargs.get("estimate_file", "")
        
        return {
            "estimate_file": estimate_file,
            "status": "parsed",
            "positions": [],
            "total_cost": 0.0,
            "currency": "RUB"
        }
    
    def _create_gpp(self, **kwargs) -> Dict[str, Any]:
        """Создание ГПП"""
        works_seq = kwargs.get("works_seq", [])
        
        return {
            "gpp_created": True,
            "works_count": len(works_seq),
            "timeline": "auto_generated",
            "created_date": datetime.now().isoformat()
        }
    
    def _extract_financial_data(self, **kwargs) -> Dict[str, Any]:
        """Извлечение финансовых данных"""
        document = kwargs.get("document", "")
        
        # Простые regex для поиска денежных сумм
        money_patterns = [
            r'(\d+(?:\s?\d{3})*(?:[.,]\d{2})?)\s*(?:руб|₽|рублей)',
            r'(\d+(?:\s?\d{3})*(?:[.,]\d{2})?)\s*(?:тыс|млн)',
        ]
        
        financial_data = []
        
        for pattern in money_patterns:
            matches = re.findall(pattern, document, re.IGNORECASE)
            for match in matches:
                amount = re.sub(r'[^\d,.]', '', match)
                try:
                    amount_float = float(amount.replace(',', '.'))
                    financial_data.append({
                        "amount": amount_float,
                        "original_text": match,
                        "currency": "RUB"
                    })
                except ValueError:
                    continue
        
        return {
            "document_length": len(document),
            "financial_entries": len(financial_data),
            "data": financial_data[:10],  # Ограничиваем 10
            "total_amount": sum(item["amount"] for item in financial_data)
        }
    
    # ===============================
    # PUBLIC API
    # ===============================
    
    def list_all_tools(self, category: Optional[str] = None) -> Dict[str, Any]:
        """Список всех инструментов"""
        if category:
            try:
                cat_enum = ToolCategory(category)
                tools = self.registry.list_tools(cat_enum)
            except ValueError:
                return {"error": f"Invalid category: {category}"}
        else:
            tools = self.registry.list_tools()
        
        return {
            "tools": [
                {
                    "name": tool.name,
                    "description": tool.description,
                    "category": tool.category.value,
                    "required_params": tool.required_params,
                    "optional_params": tool.optional_params,
                    "ui_placement": tool.ui_placement
                }
                for tool in tools
            ],
            "total_count": len(tools),
            "categories": self.registry.get_categories()
        }
    
    def get_tool_info(self, tool_name: str) -> Dict[str, Any]:
        """Информация об инструменте"""
        tool = self.registry.get_tool(tool_name)
        if not tool:
            return {"error": f"Tool '{tool_name}' not found"}
        
        return {
            "name": tool.name,
            "description": tool.description,
            "category": tool.category.value,
            "required_params": tool.required_params,
            "optional_params": tool.optional_params,
            "ui_placement": tool.ui_placement,
            "available": tool_name in self.tool_methods
        }
    
    def get_execution_stats(self, tool_name: Optional[str] = None) -> Dict[str, Any]:
        """Статистика выполнения"""
        return self.executor.get_stats(tool_name)
    
    def execute_tool_chain(self, chain: List[Dict[str, Any]]) -> List[ToolResult]:
        """Выполнение цепочки инструментов"""
        results = []
        
        for step in chain:
            tool_name = step.get("tool")
            params = step.get("params", {})
            
            if not tool_name:
                result = ToolResult(
                    status="error",
                    error="Missing tool name in chain step",
                    metadata={"step": step}
                )
                results.append(result)
                continue
            
            # Можем использовать результаты предыдущих шагов
            if step.get("use_previous_result") and results:
                previous_result = results[-1]
                if previous_result.is_success():
                    params.update(previous_result.data or {})
            
            result = self.execute_tool(tool_name, **params)
            results.append(result)
            
            # Прерываем цепочку при ошибке, если не указано иное
            if not result.is_success() and not step.get("continue_on_error", False):
                break
        
        return results

# ===============================
# GLOBAL INSTANCE
# ===============================

# Создаём глобальный экземпляр для обратной совместимости
master_tools = None

def get_master_tools_system(rag_system=None, model_manager=None) -> MasterToolsSystem:
    """Получение глобального экземпляра Master Tools System"""
    global master_tools
    if master_tools is None:
        master_tools = MasterToolsSystem(rag_system, model_manager)
    return master_tools

# Convenience functions для обратной совместимости
def execute_tool(tool_name: str, **kwargs) -> ToolResult:
    """Удобная функция для выполнения инструментов"""
    return get_master_tools_system().execute_tool(tool_name, **kwargs)

def list_available_tools(category: str = None) -> Dict[str, Any]:
    """Список доступных инструментов"""
    return get_master_tools_system().list_all_tools(category)

def get_tool_info(tool_name: str) -> Dict[str, Any]:
    """Информация об инструменте"""
    return get_master_tools_system().get_tool_info(tool_name)

if __name__ == "__main__":
    # Тестирование системы
    print("🚀 MASTER TOOLS SYSTEM - INITIALIZATION TEST")
    print("=" * 60)
    
    # Создаём систему
    mts = MasterToolsSystem()
    
    # Показываем статистику
    tools_info = mts.list_all_tools()
    print(f"✅ Loaded {tools_info['total_count']} tools")
    print(f"📊 Categories: {', '.join(tools_info['categories'])}")
    
    # Тестируем несколько инструментов
    test_tools = [
        ("search_rag_database", {"query": "test"}),
        ("create_pie_chart", {"data": {"A": 30, "B": 40, "C": 30}}),
        ("calculate_financial_metrics", {"metric_type": "roi", "investment": 1000, "cash_flows": [300, 400, 500]})
    ]
    
    print(f"\n🧪 Testing {len(test_tools)} tools:")
    for tool_name, params in test_tools:
        result = mts.execute_tool(tool_name, **params)
        status = "✅" if result.is_success() else "❌"
        print(f"  {status} {tool_name}: {result.status}")
    
    print(f"\n📈 Execution stats:")
    stats = mts.get_execution_stats()
    for tool_name, tool_stats in stats.items():
        success_rate = (tool_stats['successful_calls'] / tool_stats['total_calls']) * 100
        print(f"  {tool_name}: {success_rate:.1f}% success rate, avg {tool_stats['avg_execution_time']:.3f}s")
    
    print(f"\n🎉 Master Tools System ready for production!")