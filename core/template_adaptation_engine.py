#!/usr/bin/env python3
"""
Template Adaptation Engine
Движок адаптации шаблонов документов под конкретные компании и проекты
"""

import os
import re
import json
import shutil
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, field
from datetime import datetime, date
from enum import Enum
import logging

# Импорты для работы с документами
try:
    import docx
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    docx = None

try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError:
    canvas = None

logger = logging.getLogger(__name__)

class AdaptationType(Enum):
    COMPANY_DATA = "company_data"
    PROJECT_DATA = "project_data"
    CUSTOM_FIELDS = "custom_fields"
    FORMATTING = "formatting"
    LOCALIZATION = "localization"

class PlaceholderType(Enum):
    TEXT = "text"
    NUMBER = "number"
    DATE = "date"
    EMAIL = "email"
    PHONE = "phone"
    ADDRESS = "address"
    MONEY = "money"
    SIGNATURE = "signature"

@dataclass
class CompanyInfo:
    name: str = ""
    full_name: str = ""
    address: str = ""
    postal_code: str = ""
    city: str = ""
    country: str = ""
    phone: str = ""
    email: str = ""
    website: str = ""
    inn: str = ""
    kpp: str = ""
    ogrn: str = ""
    director: str = ""
    director_position: str = "Генеральный директор"
    accountant: str = ""
    bank_name: str = ""
    bank_account: str = ""
    bank_bik: str = ""
    bank_correspondent_account: str = ""
    
    def to_dict(self) -> Dict[str, str]:
        """Преобразование в словарь для замены плейсхолдеров"""
        return {
            'COMPANY_NAME': self.name,
            'COMPANY_FULL_NAME': self.full_name or self.name,
            'COMPANY_ADDRESS': self.address,
            'COMPANY_POSTAL_CODE': self.postal_code,
            'COMPANY_CITY': self.city,
            'COMPANY_COUNTRY': self.country,
            'COMPANY_PHONE': self.phone,
            'COMPANY_EMAIL': self.email,
            'COMPANY_WEBSITE': self.website,
            'COMPANY_INN': self.inn,
            'COMPANY_KPP': self.kpp,
            'COMPANY_OGRN': self.ogrn,
            'DIRECTOR_NAME': self.director,
            'DIRECTOR_POSITION': self.director_position,
            'ACCOUNTANT_NAME': self.accountant,
            'BANK_NAME': self.bank_name,
            'BANK_ACCOUNT': self.bank_account,
            'BANK_BIK': self.bank_bik,
            'BANK_CORRESPONDENT_ACCOUNT': self.bank_correspondent_account
        }

@dataclass
class ProjectInfo:
    name: str = ""
    description: str = ""
    type: str = ""
    location: str = ""
    address: str = ""
    budget: float = 0.0
    currency: str = "RUB"
    start_date: Optional[date] = None
    end_date: Optional[date] = None
    manager: str = ""
    client: str = ""
    contract_number: str = ""
    
    def to_dict(self) -> Dict[str, str]:
        """Преобразование в словарь для замены плейсхолдеров"""
        return {
            'PROJECT_NAME': self.name,
            'PROJECT_DESCRIPTION': self.description,
            'PROJECT_TYPE': self.type,
            'PROJECT_LOCATION': self.location,
            'PROJECT_ADDRESS': self.address,
            'PROJECT_BUDGET': f"{self.budget:,.2f}" if self.budget else "",
            'PROJECT_CURRENCY': self.currency,
            'PROJECT_START_DATE': self.start_date.strftime('%d.%m.%Y') if self.start_date else "",
            'PROJECT_END_DATE': self.end_date.strftime('%d.%m.%Y') if self.end_date else "",
            'PROJECT_MANAGER': self.manager,
            'PROJECT_CLIENT': self.client,
            'CONTRACT_NUMBER': self.contract_number
        }

@dataclass
class AdaptationRule:
    placeholder_pattern: str
    replacement_value: str
    adaptation_type: AdaptationType
    placeholder_type: PlaceholderType = PlaceholderType.TEXT
    required: bool = True
    validation_pattern: Optional[str] = None
    format_function: Optional[str] = None

@dataclass
class AdaptationResult:
    success: bool
    adapted_file_path: str = ""
    original_file_path: str = ""
    replacements_made: int = 0
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    processing_time: float = 0.0c
lass TemplateAdaptationEngine:
    """Движок адаптации шаблонов документов"""
    
    def __init__(self):
        self.templates_dir = Path("I:/docs/templates")
        self.generated_dir = Path("I:/docs_generated/templates")
        self.cache_dir = Path("I:/docs/cache/templates")
        
        # Создаем директории
        for dir_path in [self.templates_dir, self.generated_dir, self.cache_dir]:
            dir_path.mkdir(parents=True, exist_ok=True)
        
        # Стандартные паттерны плейсхолдеров
        self.placeholder_patterns = {
            # Фигурные скобки
            'curly_braces': r'\{([^}]+)\}',
            # Квадратные скобки
            'square_brackets': r'\[([^\]]+)\]',
            # Угловые скобки
            'angle_brackets': r'<<([^>]+)>>',
            # Процентные знаки
            'percent_signs': r'%([^%]+)%',
            # Подчеркивания
            'underscores': r'_{3,}',
            # Точки
            'dots': r'\.{3,}'
        }
        
        # Форматтеры для различных типов данных
        self.formatters = {
            PlaceholderType.DATE: self._format_date,
            PlaceholderType.MONEY: self._format_money,
            PlaceholderType.PHONE: self._format_phone,
            PlaceholderType.EMAIL: self._format_email,
            PlaceholderType.NUMBER: self._format_number
        }
        
        # Валидаторы
        self.validators = {
            PlaceholderType.EMAIL: r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$',
            PlaceholderType.PHONE: r'^[\+]?[1-9][\d]{0,15}$',
            PlaceholderType.DATE: r'^\d{2}\.\d{2}\.\d{4}$'
        }
        
        # Стандартные замены
        self.standard_replacements = {
            'CURRENT_DATE': datetime.now().strftime('%d.%m.%Y'),
            'CURRENT_YEAR': str(datetime.now().year),
            'CURRENT_MONTH': datetime.now().strftime('%m'),
            'CURRENT_DAY': datetime.now().strftime('%d')
        }
    
    def adapt_template(self, template_path: str, company_info: CompanyInfo = None,
                      project_info: ProjectInfo = None, custom_fields: Dict[str, str] = None,
                      output_filename: str = None) -> AdaptationResult:
        """
        Адаптация шаблона под конкретные данные
        
        Args:
            template_path: Путь к исходному шаблону
            company_info: Информация о компании
            project_info: Информация о проекте
            custom_fields: Дополнительные поля для замены
            output_filename: Имя выходного файла
        """
        start_time = datetime.now()
        
        try:
            template_file = Path(template_path)
            if not template_file.exists():
                return AdaptationResult(
                    success=False,
                    errors=[f"Template file not found: {template_path}"]
                )
            
            # Определяем выходной файл
            if not output_filename:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                output_filename = f"{template_file.stem}_adapted_{timestamp}{template_file.suffix}"
            
            output_path = self.generated_dir / output_filename
            
            # Собираем все замены
            replacements = self._collect_replacements(company_info, project_info, custom_fields)
            
            # Выполняем адаптацию в зависимости от формата файла
            if template_file.suffix.lower() == '.docx':
                result = self._adapt_docx_template(template_file, output_path, replacements)
            elif template_file.suffix.lower() in ['.txt', '.md']:
                result = self._adapt_text_template(template_file, output_path, replacements)
            elif template_file.suffix.lower() == '.html':
                result = self._adapt_html_template(template_file, output_path, replacements)
            else:
                # Для других форматов пытаемся как текстовый файл
                result = self._adapt_text_template(template_file, output_path, replacements)
            
            result.original_file_path = str(template_file)
            result.adapted_file_path = str(output_path)
            
            # Создаем метаданные адаптации
            self._create_adaptation_metadata(result, template_file, replacements)
            
        except Exception as e:
            logger.error(f"Error adapting template: {e}")
            result = AdaptationResult(
                success=False,
                errors=[str(e)]
            )
        
        result.processing_time = (datetime.now() - start_time).total_seconds()
        return result
    
    def _collect_replacements(self, company_info: CompanyInfo = None,
                            project_info: ProjectInfo = None,
                            custom_fields: Dict[str, str] = None) -> Dict[str, str]:
        """Сбор всех замен для плейсхолдеров"""
        replacements = {}
        
        # Стандартные замены
        replacements.update(self.standard_replacements)
        
        # Данные компании
        if company_info:
            replacements.update(company_info.to_dict())
        
        # Данные проекта
        if project_info:
            replacements.update(project_info.to_dict())
        
        # Пользовательские поля
        if custom_fields:
            replacements.update(custom_fields)
        
        return replacements
    
    def _adapt_docx_template(self, template_path: Path, output_path: Path,
                           replacements: Dict[str, str]) -> AdaptationResult:
        """Адаптация DOCX шаблона"""
        if not docx:
            return AdaptationResult(
                success=False,
                errors=["python-docx library not available"]
            )
        
        try:
            # Открываем документ
            doc = docx.Document(template_path)
            replacements_made = 0
            
            # Заменяем в параграфах
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                new_text = self._replace_placeholders(original_text, replacements)
                if new_text != original_text:
                    paragraph.text = new_text
                    replacements_made += 1
            
            # Заменяем в таблицах
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text
                            new_text = self._replace_placeholders(original_text, replacements)
                            if new_text != original_text:
                                paragraph.text = new_text
                                replacements_made += 1
            
            # Заменяем в заголовках и колонтитулах
            for section in doc.sections:
                # Заголовок
                if section.header:
                    for paragraph in section.header.paragraphs:
                        original_text = paragraph.text
                        new_text = self._replace_placeholders(original_text, replacements)
                        if new_text != original_text:
                            paragraph.text = new_text
                            replacements_made += 1
                
                # Колонтитул
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        original_text = paragraph.text
                        new_text = self._replace_placeholders(original_text, replacements)
                        if new_text != original_text:
                            paragraph.text = new_text
                            replacements_made += 1
            
            # Сохраняем документ
            doc.save(output_path)
            
            return AdaptationResult(
                success=True,
                replacements_made=replacements_made
            )
            
        except Exception as e:
            return AdaptationResult(
                success=False,
                errors=[f"Error processing DOCX: {e}"]
            )
    
    def _adapt_text_template(self, template_path: Path, output_path: Path,
                           replacements: Dict[str, str]) -> AdaptationResult:
        """Адаптация текстового шаблона"""
        try:
            # Читаем файл
            content = self._read_text_file(template_path)
            
            # Выполняем замены
            original_content = content
            adapted_content = self._replace_placeholders(content, replacements)
            
            # Подсчитываем количество замен
            replacements_made = self._count_replacements(original_content, adapted_content, replacements)
            
            # Сохраняем адаптированный файл
            self._write_text_file(output_path, adapted_content)
            
            return AdaptationResult(
                success=True,
                replacements_made=replacements_made
            )
            
        except Exception as e:
            return AdaptationResult(
                success=False,
                errors=[f"Error processing text file: {e}"]
            )
    
    def _adapt_html_template(self, template_path: Path, output_path: Path,
                           replacements: Dict[str, str]) -> AdaptationResult:
        """Адаптация HTML шаблона"""
        try:
            # Читаем HTML файл
            content = self._read_text_file(template_path)
            
            # Выполняем замены
            original_content = content
            adapted_content = self._replace_placeholders(content, replacements)
            
            # Подсчитываем количество замен
            replacements_made = self._count_replacements(original_content, adapted_content, replacements)
            
            # Сохраняем адаптированный файл
            self._write_text_file(output_path, adapted_content)
            
            return AdaptationResult(
                success=True,
                replacements_made=replacements_made
            )
            
        except Exception as e:
            return AdaptationResult(
                success=False,
                errors=[f"Error processing HTML file: {e}"]
            )    
 
   def _replace_placeholders(self, text: str, replacements: Dict[str, str]) -> str:
        """Замена плейсхолдеров в тексте"""
        result = text
        
        # Проходим по всем паттернам плейсхолдеров
        for pattern_name, pattern in self.placeholder_patterns.items():
            if pattern_name in ['underscores', 'dots']:
                # Для подчеркиваний и точек используем специальную логику
                continue
            
            # Ищем плейсхолдеры по паттерну
            matches = re.finditer(pattern, result, re.IGNORECASE)
            
            for match in reversed(list(matches)):  # Обрабатываем в обратном порядке
                placeholder_full = match.group(0)
                placeholder_name = match.group(1) if len(match.groups()) > 0 else placeholder_full
                
                # Ищем замену
                replacement = self._find_replacement(placeholder_name, replacements)
                
                if replacement is not None:
                    # Применяем форматирование если нужно
                    formatted_replacement = self._apply_formatting(placeholder_name, replacement)
                    
                    # Заменяем плейсхолдер
                    result = result[:match.start()] + formatted_replacement + result[match.end():]
        
        # Обрабатываем простые замены без скобок
        for key, value in replacements.items():
            # Ищем точные совпадения ключей
            result = re.sub(rf'\b{re.escape(key)}\b', value, result, flags=re.IGNORECASE)
        
        return result
    
    def _find_replacement(self, placeholder_name: str, replacements: Dict[str, str]) -> Optional[str]:
        """Поиск замены для плейсхолдера"""
        # Нормализуем имя плейсхолдера
        normalized_name = placeholder_name.upper().strip()
        
        # Прямое совпадение
        if normalized_name in replacements:
            return replacements[normalized_name]
        
        # Поиск по частичному совпадению
        for key, value in replacements.items():
            if key.upper() in normalized_name or normalized_name in key.upper():
                return value
        
        # Поиск по синонимам
        synonyms = {
            'ФИО': ['FULL_NAME', 'NAME', 'DIRECTOR_NAME'],
            'ДАТА': ['DATE', 'CURRENT_DATE'],
            'СУММА': ['AMOUNT', 'BUDGET', 'MONEY'],
            'АДРЕС': ['ADDRESS', 'COMPANY_ADDRESS'],
            'ТЕЛЕФОН': ['PHONE', 'COMPANY_PHONE'],
            'EMAIL': ['COMPANY_EMAIL', 'MAIL'],
            'НАЗВАНИЕ': ['NAME', 'COMPANY_NAME'],
            'ДИРЕКТОР': ['DIRECTOR', 'DIRECTOR_NAME'],
            'ГОРОД': ['CITY', 'COMPANY_CITY']
        }
        
        for synonym_key, synonym_list in synonyms.items():
            if synonym_key in normalized_name:
                for synonym in synonym_list:
                    if synonym in replacements:
                        return replacements[synonym]
        
        return None
    
    def _apply_formatting(self, placeholder_name: str, value: str) -> str:
        """Применение форматирования к значению"""
        placeholder_lower = placeholder_name.lower()
        
        # Определяем тип плейсхолдера и применяем соответствующее форматирование
        if any(word in placeholder_lower for word in ['дата', 'date']):
            return self._format_date(value)
        elif any(word in placeholder_lower for word in ['сумма', 'money', 'budget', 'amount']):
            return self._format_money(value)
        elif any(word in placeholder_lower for word in ['телефон', 'phone']):
            return self._format_phone(value)
        elif any(word in placeholder_lower for word in ['email', 'почта']):
            return self._format_email(value)
        
        return value
    
    def _format_date(self, value: str) -> str:
        """Форматирование даты"""
        if not value:
            return ""
        
        # Если уже в правильном формате
        if re.match(r'^\d{2}\.\d{2}\.\d{4}$', value):
            return value
        
        # Пытаемся распарсить различные форматы
        date_patterns = [
            r'(\d{4})-(\d{2})-(\d{2})',  # YYYY-MM-DD
            r'(\d{2})/(\d{2})/(\d{4})',  # DD/MM/YYYY
            r'(\d{2})-(\d{2})-(\d{4})'   # DD-MM-YYYY
        ]
        
        for pattern in date_patterns:
            match = re.match(pattern, value)
            if match:
                if len(match.group(1)) == 4:  # Год в начале
                    year, month, day = match.groups()
                else:  # День в начале
                    day, month, year = match.groups()
                return f"{day}.{month}.{year}"
        
        return value
    
    def _format_money(self, value: str) -> str:
        """Форматирование денежной суммы"""
        if not value:
            return ""
        
        try:
            # Удаляем все кроме цифр и точки/запятой
            clean_value = re.sub(r'[^\d.,]', '', value)
            clean_value = clean_value.replace(',', '.')
            
            amount = float(clean_value)
            
            # Форматируем с разделителями тысяч
            return f"{amount:,.2f}".replace(',', ' ')
        except ValueError:
            return value
    
    def _format_phone(self, value: str) -> str:
        """Форматирование номера телефона"""
        if not value:
            return ""
        
        # Удаляем все кроме цифр и +
        clean_phone = re.sub(r'[^\d+]', '', value)
        
        # Российский номер
        if clean_phone.startswith('8') and len(clean_phone) == 11:
            clean_phone = '+7' + clean_phone[1:]
        elif clean_phone.startswith('7') and len(clean_phone) == 11:
            clean_phone = '+' + clean_phone
        
        # Форматируем российский номер
        if clean_phone.startswith('+7') and len(clean_phone) == 12:
            return f"+7 ({clean_phone[2:5]}) {clean_phone[5:8]}-{clean_phone[8:10]}-{clean_phone[10:12]}"
        
        return value
    
    def _format_email(self, value: str) -> str:
        """Форматирование email"""
        if not value:
            return ""
        
        # Простая валидация email
        if re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', value):
            return value.lower()
        
        return value
    
    def _format_number(self, value: str) -> str:
        """Форматирование числа"""
        if not value:
            return ""
        
        try:
            # Пытаемся преобразовать в число
            if '.' in value or ',' in value:
                num = float(value.replace(',', '.'))
                return f"{num:g}"  # Убирает лишние нули
            else:
                num = int(value)
                return str(num)
        except ValueError:
            return value
    
    def _count_replacements(self, original: str, adapted: str, replacements: Dict[str, str]) -> int:
        """Подсчет количества выполненных замен"""
        count = 0
        
        # Подсчитываем замены для каждого ключа
        for key, value in replacements.items():
            if value and value in adapted:
                # Подсчитываем количество вхождений
                count += adapted.count(value) - original.count(value)
        
        return max(0, count)  # Не может быть отрицательным
    
    def _read_text_file(self, file_path: Path) -> str:
        """Чтение текстового файла с автоопределением кодировки"""
        encodings = ['utf-8', 'cp1251', 'latin1', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Could not decode file {file_path}")
    
    def _write_text_file(self, file_path: Path, content: str):
        """Запись текстового файла"""
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(content)
    
    def _create_adaptation_metadata(self, result: AdaptationResult, template_file: Path,
                                  replacements: Dict[str, str]):
        """Создание метаданных адаптации"""
        metadata = {
            'template_file': str(template_file),
            'adaptation_date': datetime.now().isoformat(),
            'replacements_count': result.replacements_made,
            'template_size': template_file.stat().st_size,
            'replacements_used': {k: v for k, v in replacements.items() if v}
        }
        
        result.metadata = metadata
        
        # Сохраняем метаданные в отдельный файл
        if result.adapted_file_path:
            metadata_file = Path(result.adapted_file_path).with_suffix('.json')
            with open(metadata_file, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, ensure_ascii=False, indent=2)
    
    def create_template_from_text(self, text_content: str, template_name: str,
                                category: str = "custom") -> str:
        """Создание шаблона из текста"""
        try:
            # Создаем файл шаблона
            template_file = self.templates_dir / f"{template_name}.txt"
            
            with open(template_file, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            # Создаем метаданные шаблона
            metadata = {
                'name': template_name,
                'category': category,
                'created_date': datetime.now().isoformat(),
                'format': 'txt',
                'description': f"Custom template: {template_name}"
            }
            
            metadata_file = template_file.with_suffix('.meta.json')
            with open(metadata_file, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, ensure_ascii=False, indent=2)
            
            return str(template_file)
            
        except Exception as e:
            logger.error(f"Error creating template: {e}")
            raise
    
    def get_template_placeholders(self, template_path: str) -> List[str]:
        """Извлечение списка плейсхолдеров из шаблона"""
        try:
            template_file = Path(template_path)
            
            if template_file.suffix.lower() == '.docx' and docx:
                content = self._extract_docx_text(template_file)
            else:
                content = self._read_text_file(template_file)
            
            placeholders = set()
            
            # Ищем плейсхолдеры по всем паттернам
            for pattern_name, pattern in self.placeholder_patterns.items():
                if pattern_name in ['underscores', 'dots']:
                    continue
                
                matches = re.finditer(pattern, content, re.IGNORECASE)
                for match in matches:
                    if len(match.groups()) > 0:
                        placeholders.add(match.group(1))
                    else:
                        placeholders.add(match.group(0))
            
            return sorted(list(placeholders))
            
        except Exception as e:
            logger.error(f"Error extracting placeholders: {e}")
            return []
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Извлечение текста из DOCX файла"""
        if not docx:
            return ""
        
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""# Фу
нкции для интеграции с унифицированной системой инструментов
def adapt_template_for_company(template_path: str, company_info: Dict[str, str],
                             project_info: Dict[str, str] = None,
                             custom_fields: Dict[str, str] = None,
                             output_filename: str = None, **kwargs) -> Dict[str, Any]:
    """
    Адаптация шаблона под компанию и проект
    
    Args:
        template_path: Путь к шаблону
        company_info: Информация о компании
        project_info: Информация о проекте (опционально)
        custom_fields: Дополнительные поля (опционально)
        output_filename: Имя выходного файла (опционально)
        **kwargs: Дополнительные параметры
    """
    try:
        engine = TemplateAdaptationEngine()
        
        # Преобразуем словари в объекты
        company = CompanyInfo(**company_info) if company_info else None
        project = ProjectInfo(**project_info) if project_info else None
        
        # Выполняем адаптацию
        result = engine.adapt_template(
            template_path=template_path,
            company_info=company,
            project_info=project,
            custom_fields=custom_fields,
            output_filename=output_filename
        )
        
        return {
            'status': 'success' if result.success else 'error',
            'adapted_file': result.adapted_file_path,
            'original_file': result.original_file_path,
            'replacements_made': result.replacements_made,
            'processing_time': result.processing_time,
            'metadata': result.metadata,
            'errors': result.errors,
            'warnings': result.warnings
        }
        
    except Exception as e:
        logger.error(f"Error in template adaptation: {e}")
        return {
            'status': 'error',
            'error': str(e),
            'template_path': template_path
        }

def create_custom_template(text_content: str, template_name: str,
                         category: str = "custom", **kwargs) -> Dict[str, Any]:
    """
    Создание пользовательского шаблона
    
    Args:
        text_content: Содержимое шаблона
        template_name: Название шаблона
        category: Категория шаблона
        **kwargs: Дополнительные параметры
    """
    try:
        engine = TemplateAdaptationEngine()
        
        template_path = engine.create_template_from_text(
            text_content=text_content,
            template_name=template_name,
            category=category
        )
        
        return {
            'status': 'success',
            'template_path': template_path,
            'template_name': template_name,
            'category': category
        }
        
    except Exception as e:
        logger.error(f"Error creating template: {e}")
        return {
            'status': 'error',
            'error': str(e),
            'template_name': template_name
        }

def get_template_placeholders_list(template_path: str, **kwargs) -> Dict[str, Any]:
    """
    Получение списка плейсхолдеров в шаблоне
    
    Args:
        template_path: Путь к шаблону
        **kwargs: Дополнительные параметры
    """
    try:
        engine = TemplateAdaptationEngine()
        placeholders = engine.get_template_placeholders(template_path)
        
        return {
            'status': 'success',
            'template_path': template_path,
            'placeholders': placeholders,
            'placeholders_count': len(placeholders)
        }
        
    except Exception as e:
        logger.error(f"Error getting placeholders: {e}")
        return {
            'status': 'error',
            'error': str(e),
            'template_path': template_path
        }

# Регистрация в унифицированной системе инструментов
TEMPLATE_ADAPTATION_TOOLS = {
    'adapt_template_for_company': {
        'function': adapt_template_for_company,
        'description': 'Адаптация шаблона документа под конкретную компанию и проект',
        'category': 'document_generation',
        'ui_placement': 'tools',
        'parameters': {
            'template_path': 'str - Путь к файлу шаблона',
            'company_info': 'dict - Информация о компании (name, address, phone, email, inn, etc.)',
            'project_info': 'dict - Информация о проекте (name, description, budget, etc.)',
            'custom_fields': 'dict - Дополнительные поля для замены',
            'output_filename': 'str - Имя выходного файла (опционально)'
        }
    },
    'create_custom_template': {
        'function': create_custom_template,
        'description': 'Создание пользовательского шаблона из текста',
        'category': 'document_generation',
        'ui_placement': 'tools',
        'parameters': {
            'text_content': 'str - Содержимое шаблона с плейсхолдерами',
            'template_name': 'str - Название шаблона',
            'category': 'str - Категория шаблона'
        }
    },
    'get_template_placeholders_list': {
        'function': get_template_placeholders_list,
        'description': 'Получение списка плейсхолдеров в шаблоне',
        'category': 'document_analysis',
        'ui_placement': 'tools',
        'parameters': {
            'template_path': 'str - Путь к файлу шаблона'
        }
    }
}

if __name__ == "__main__":
    # Тестирование движка адаптации
    engine = TemplateAdaptationEngine()
    
    # Создаем тестовый шаблон
    test_template_content = """
ДОГОВОР СТРОИТЕЛЬНОГО ПОДРЯДА №{CONTRACT_NUMBER}

г. {COMPANY_CITY}, {CURRENT_DATE}

{COMPANY_NAME}, в лице {DIRECTOR_POSITION} {DIRECTOR_NAME}, действующего на основании Устава,
именуемое в дальнейшем "Заказчик", с одной стороны, и 

{PROJECT_CLIENT}, именуемый в дальнейшем "Подрядчик", с другой стороны,
заключили настоящий договор о нижеследующем:

1. ПРЕДМЕТ ДОГОВОРА

1.1. Подрядчик обязуется выполнить строительные работы по объекту:
"{PROJECT_NAME}" по адресу: {PROJECT_ADDRESS}

1.2. Общая стоимость работ составляет: {PROJECT_BUDGET} {PROJECT_CURRENCY}

2. СРОКИ ВЫПОЛНЕНИЯ РАБОТ

2.1. Начало работ: {PROJECT_START_DATE}
2.2. Окончание работ: {PROJECT_END_DATE}

3. РЕКВИЗИТЫ ЗАКАЗЧИКА

Наименование: {COMPANY_FULL_NAME}
Адрес: {COMPANY_ADDRESS}
ИНН: {COMPANY_INN}
КПП: {COMPANY_KPP}
Телефон: {COMPANY_PHONE}
E-mail: {COMPANY_EMAIL}

Подписи сторон:

Заказчик: _________________ {DIRECTOR_NAME}

Подрядчик: ________________ {PROJECT_CLIENT}
"""
    
    # Создаем тестовый шаблон
    template_path = engine.create_template_from_text(
        test_template_content,
        "test_contract_template",
        "contracts"
    )
    
    print(f"Created test template: {template_path}")
    
    # Тестовые данные компании
    company_info = CompanyInfo(
        name="ООО 'СтройИнвест'",
        full_name="Общество с ограниченной ответственностью 'СтройИнвест'",
        address="123456, г. Москва, ул. Строительная, д. 10",
        city="Москва",
        phone="+7 (495) 123-45-67",
        email="info@stroyinvest.ru",
        inn="7701234567",
        kpp="770101001",
        director="Иванов Иван Иванович"
    )
    
    # Тестовые данные проекта
    project_info = ProjectInfo(
        name="Строительство жилого дома",
        description="Строительство 3-этажного жилого дома",
        address="г. Москва, ул. Новая, участок 15",
        budget=5000000.0,
        currency="RUB",
        start_date=date(2024, 3, 1),
        end_date=date(2024, 12, 31),
        client="ИП Петров П.П.",
        contract_number="СД-2024-001"
    )
    
    # Дополнительные поля
    custom_fields = {
        "CONTRACT_NUMBER": "СД-2024-001"
    }
    
    # Выполняем адаптацию
    print("\nAdapting template...")
    result = engine.adapt_template(
        template_path=template_path,
        company_info=company_info,
        project_info=project_info,
        custom_fields=custom_fields,
        output_filename="adapted_contract.txt"
    )
    
    print(f"Adaptation result:")
    print(f"  Success: {result.success}")
    print(f"  Adapted file: {result.adapted_file_path}")
    print(f"  Replacements made: {result.replacements_made}")
    print(f"  Processing time: {result.processing_time:.3f} sec")
    
    if result.errors:
        print(f"  Errors: {result.errors}")
    
    if result.warnings:
        print(f"  Warnings: {result.warnings}")
    
    # Показываем содержимое адаптированного файла
    if result.success and result.adapted_file_path:
        print(f"\nAdapted content preview:")
        with open(result.adapted_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            print(content[:500] + "..." if len(content) > 500 else content)
    
    # Тестируем извлечение плейсхолдеров
    print(f"\nPlaceholders in template:")
    placeholders = engine.get_template_placeholders(template_path)
    for placeholder in placeholders:
        print(f"  - {placeholder}")
    
    print(f"\nTotal placeholders found: {len(placeholders)}")